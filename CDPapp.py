# aydi: UTAS CDP Builder — PATCHED
# Update CDPapp.py
# Streamlit app: UTAS CDP Builder — daily CDP authoring with docxtpl render.

import os
import io
from pathlib import Path
import inspect

import streamlit as st
import pandas as pd
import yaml

from docxtpl import DocxTemplate, RichText
from docx.shared import Inches, Length
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.shared import OxmlElement, qn
import json
import requests
from datetime import datetime, date
# --- Sign-off imports ---
import time, json, hashlib, secrets
from pathlib import Path
from PIL import Image
import numpy as np
import hashlib
import tempfile
from pathlib import Path
import streamlit.components.v1 as components
#---notifications by gmail----
import smtplib
from email.message import EmailMessage

def _smtp_cfg() -> dict:
    return dict(st.secrets.get("smtp", {})) if hasattr(st, "secrets") else {}

def _send_email(to_email: str, subject: str, body: str) -> tuple[bool, str]:
    cfg = _smtp_cfg()
    if not cfg:
        return False, "SMTP not configured (missing [smtp] in secrets)."

    host = cfg.get("HOST")
    port = int(cfg.get("PORT", 587))
    user = cfg.get("USERNAME")
    pw   = cfg.get("PASSWORD")

    if not (host and user and pw):
        return False, "SMTP secrets incomplete (HOST/USERNAME/PASSWORD)."

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = cfg.get("FROM", user)
    msg["To"] = to_email
    msg.set_content(body)

    try:
        with smtplib.SMTP(host, port, timeout=20) as server:
            server.starttls()
            server.login(user, pw)
            server.send_message(msg)
        return True, ""
    except Exception as e:
        return False, str(e)

def _email_signoff_request(*, token: str, to_name: str, to_email: str, meta: dict) -> tuple[bool, str]:
    cfg = _smtp_cfg()

    # Optional “test mode”: route all emails to a single inbox
    forced_to = (cfg.get("TEST_TO") or "").strip()
    actual_to = forced_to or (to_email or "").strip()
    if not actual_to:
        return False, "No recipient email (and no smtp.TEST_TO configured)."

    link = _sign_link_for(token)  # uses APP_BASE_URL

    course_code  = (meta.get("course_code") or "").strip()
    course_title = (meta.get("course_title") or "").strip()
    ay           = (meta.get("academic_year") or "").strip()
    sem          = (meta.get("semester") or "").strip()
    sections     = (meta.get("sections") or "").strip()
    role         = meta.get("row_type", "sign-off")

    subject = f"UTAS CDP Builder: Sign-off requested ({course_code})"

    body = f"""Hello {to_name or "Colleague"},

You have a pending {role} task in UTAS CDP Builder.

Course: {course_code} — {course_title}
Academic Year: {ay}
Semester: {sem}
Sections: {sections}

Open and sign here (unique link):
{link}

If you were not expecting this, please ignore this email.

— UTAS CDP Builder
"""
    return _send_email(actual_to, subject, body)
#--- end of notifications helpers---#

from audit_handouts import (
    build_corpus_from_uploads,
    build_corpus,
    parse_pdf,
    parse_pptx,
    prepare_llm_payload,
    run_handout_audit,
    render_audit_summary_docx,
    render_course_audit_form_docx,
)

try:
    # draw canvas
    from streamlit_drawable_canvas import st_canvas
except Exception:
    st.warning("`streamlit-drawable-canvas` not installed. Add it to requirements.txt.")
# ==== Digital sign-off storage helpers (MUST be defined before use) ====
from pathlib import Path
import json, os, time, hashlib, secrets

DATA_DIR = Path("data"); DATA_DIR.mkdir(parents=True, exist_ok=True)
SIG_DIR  = DATA_DIR / "signatures"; SIG_DIR.mkdir(parents=True, exist_ok=True)
import re
import hashlib

HANDOUTS_DIR = DATA_DIR / "handouts_uploads"
HANDOUTS_DIR.mkdir(parents=True, exist_ok=True)

def _save_uploaded_files_pairs(uploaded_files, prefix: str) -> list[tuple[Path, str]]:
    """
    Save UploadedFile objects to HANDOUTS_DIR with stable hashed names,
    but return (saved_path, original_filename) pairs so we can show humans
    the original names in audit evidence.
    """
    out: list[tuple[Path, str]] = []
    if not uploaded_files:
        return out

    for uf in uploaded_files:
        b = uf.getvalue()
        h = hashlib.sha256(b).hexdigest()[:16]
        orig_name = Path(uf.name).name
        ext = Path(orig_name).suffix.lower()
        stem = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(orig_name).stem)[:60] or "file"
        saved_path = HANDOUTS_DIR / f"{prefix}_{stem}_{h}{ext}"
        saved_path.write_bytes(b)
        out.append((saved_path, orig_name))

    return out

def _save_uploaded_files(uploaded_files, prefix: str) -> list[Path]:
    """
    Backward-compatible wrapper returning only paths.
    """
    return [p for (p, _orig) in _save_uploaded_files_pairs(uploaded_files, prefix)]


TOK_FILE = DATA_DIR / "sign_tokens.json"    # issued tokens
REC_FILE = DATA_DIR / "sign_records.json"   # persisted signatures
LOG_FILE_SIGN = DATA_DIR / "signoff_log.jsonl"  # audit log
AI_LOG_FILE   = DATA_DIR / "ai_review_logs.jsonl"
SIGN_LOG_FILE = DATA_DIR / "signoff_log.jsonl"   # you already have LOG_FILE_SIGN -> keep that or unify names

if not TOK_FILE.exists():
    TOK_FILE.write_text("{}", encoding="utf-8")
# --- tiny helpers for one-shot buttons ---
def _busy(key: str) -> bool:
    return bool(st.session_state.get(key))

def _set_busy(key: str, val: bool = True) -> None:
    st.session_state[key] = val

def _json_load(path: Path, default=None):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {} if default is None else default

def _json_save(path: Path, obj):
    try:
        path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

def _append_sign_log(rec: dict):
    try:
        with LOG_FILE_SIGN.open("a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        pass

def _draft_id():
    d = st.session_state.get("draft", {})
    course = d.get("course", {}); doc = d.get("doc", {})
    key = "|".join([
        str(course.get("course_code","")).strip(),
        str(doc.get("academic_year","")).strip(),
        str(doc.get("semester","")).strip(),
    ])
    return hashlib.sha256(key.encode("utf-8")).hexdigest()[:16]

def _stable_json(obj) -> str:
    """Stable serialization so hashes don't change due to spacing/key order."""
    return json.dumps(obj, ensure_ascii=False, sort_keys=True, separators=(",", ":"))

def _draft_rev_for_id(draft_id: str) -> str:
    """
    Revision hash for the *current* saved snapshot of this draft.
    If no snapshot exists, hash the current bundle dict.
    """
    snap = _load_snapshot_if_any(draft_id)
    if snap is None:
        snap = _current_draft_bundle_dict() or {}
        snap["_owner_uid"] = st.session_state.get("user_code") or ""
    return hashlib.sha256(_stable_json(snap).encode("utf-8")).hexdigest()[:16]

def _close_token(tok: str, note: str = "") -> None:
    """Mark a token used/closed and optionally attach a note."""
    toks = _json_load(TOK_FILE, {})
    if tok in toks:
        toks[tok]["used_at"] = int(time.time())
        if note:
            toks[tok]["note"] = note
        _json_save(TOK_FILE, toks)

def _expire_stale_tokens_for_draft(draft_id: str) -> int:
    """
    Auto-close any open tokens for this draft whose draft_rev doesn't match
    the current snapshot rev. Prevents 'stuck' pending requests.
    """
    toks = _read_tokens()
    cur = _draft_rev_for_id(draft_id)
    now = int(time.time())
    changed = False
    count = 0

    for tok, info in toks.items():
        if info.get("draft_id") != draft_id:
            continue
        if info.get("used_at"):
            continue
        old = (info.get("draft_rev") or "").strip()
        if old and old != cur:
            info["used_at"] = now
            info["note"] = "stale_revision"
            changed = True
            count += 1

    if changed:
        _json_save(TOK_FILE, toks)
    return count

def _find_open_token(draft_id: str, row_type: str, row_index: int, name_or_email: str):
    """Return (token, info) if an UN-USED token already exists for this exact target."""
    tgt = _normalize(name_or_email)
    for tok, info in _read_tokens().items():
        if info.get("draft_id") != draft_id: 
            continue
        if info.get("row_type") != row_type: 
            continue
        if int(info.get("row_index", 0)) != int(row_index):
            continue
        if info.get("used_at"):  # already closed/used/voided
            continue
        if _normalize(info.get("email")) == tgt or _normalize(info.get("name")) == tgt:
            return tok, info
    return None, None

def _void_tokens_and_signatures_for_draft(draft_id: str, reason: str = "voided_on_rejection"):
    """Close ALL still-open tokens for this draft and wipe signature records so re-signing is required."""
    toks = _read_tokens()
    now  = int(time.time())
    changed = False
    for tok, info in toks.items():
        if info.get("draft_id") == draft_id and not info.get("used_at"):
            info["used_at"] = now
            info["note"] = reason
            changed = True
    if changed:
        _json_save(TOK_FILE, toks)

    # Clear collected signatures (keep audit trail in your JSONL)
    rec = _read_records()
    if draft_id in rec:
        rec[draft_id] = {"prepared": {}, "approved": {}}
        _json_save(REC_FILE, rec)

def _issue_sign_token(target: dict) -> str:
    tok = secrets.token_urlsafe(24)
    toks = _json_load(TOK_FILE, {})
    toks[tok] = {**target, "issued_at": int(time.time()), "used_at": None}
    _json_save(TOK_FILE, toks)
    return tok

def _mark_token_used(tok: str):
    toks = _json_load(TOK_FILE, {})
    if tok in toks:
        toks[tok]["used_at"] = int(time.time())
        _json_save(TOK_FILE, toks)
def _close_open_tokens_for_row(draft_id: str, row_type: str, row_index: int,
                              note: str = "cancelled_by_creator",
                              closed_by: str = "",
                              reason: str = "") -> list[str]:
    """
    Close ANY open tokens for (draft_id, row_type, row_index).
    Returns list of tokens that were closed.
    """
    toks = _json_load(TOK_FILE, {})
    now = int(time.time())
    closed = []
    changed = False

    for tok, info in toks.items():
        if info.get("draft_id") != draft_id:
            continue
        if info.get("row_type") != row_type:
            continue
        if int(info.get("row_index", 0)) != int(row_index):
            continue
        if info.get("used_at"):
            continue

        info["used_at"] = now
        info["note"] = note
        if closed_by:
            info["closed_by"] = closed_by
        if reason:
            info["close_reason"] = reason

        closed.append(tok)
        changed = True

    if changed:
        _json_save(TOK_FILE, toks)

    return closed


def _clear_signature_record(draft_id: str, row_type: str, row_index: int,
                           delete_file: bool = False) -> bool:
    """
    Remove signature record so re-signing is required.
    If delete_file=True, tries to remove the PNG file too.
    """
    rec_all = _json_load(REC_FILE, {})
    if draft_id not in rec_all:
        return False

    changed = False

    if row_type == "prepared":
        slot = str(row_index)
        old = (rec_all.get(draft_id, {}).get("prepared", {}) or {}).get(slot)
        if old:
            sig_path = old.get("signature_path")
            rec_all[draft_id]["prepared"].pop(slot, None)
            changed = True
            if delete_file and sig_path:
                try:
                    Path(sig_path).unlink(missing_ok=True)
                except Exception:
                    pass

    elif row_type == "approved":
        # approved always stored under "0"
        old = (rec_all.get(draft_id, {}).get("approved", {}) or {}).get("0")
        if old:
            sig_path = old.get("signature_path")
            rec_all[draft_id]["approved"].pop("0", None)
            changed = True
            if delete_file and sig_path:
                try:
                    Path(sig_path).unlink(missing_ok=True)
                except Exception:
                    pass

    if changed:
        _json_save(REC_FILE, rec_all)
    return changed


def _log_cancel_event(event: str, draft_id: str, row_type: str, row_index: int,
                      name: str = "", email: str = "", sections: str = "",
                      tokens: list[str] | None = None, reason: str = "") -> None:
    d = st.session_state.get("draft", {}) or {}
    _append_sign_log({
        "ts": int(time.time()),
        "event": event,  # e.g., request_cancelled / signature_cleared / draft_voided_by_creator
        "draft_id": draft_id,
        "row_type": row_type,
        "row_index": int(row_index),
        "name": name,
        "email": email,
        "sections": sections,
        "course_code": (d.get("course", {}) or {}).get("course_code", ""),
        "course_title": (d.get("course", {}) or {}).get("course_title", ""),
        "academic_year": (d.get("doc", {}) or {}).get("academic_year", ""),
        "semester": (d.get("doc", {}) or {}).get("semester", ""),
        "tokens_closed": tokens or [],
        "reason": (reason or "").strip(),
        "by_uid": st.session_state.get("user_code", ""),
    })

def _store_signature_record(draft_id: str, row_type: str, row_index: int, signer_name: str, sig_path: str):
    rec = _json_load(REC_FILE, {})
    rec.setdefault(draft_id, {"prepared": {}, "approved": {}})
    if row_type == "prepared":
        rec[draft_id]["prepared"][str(row_index)] = {"name": signer_name, "signature_path": sig_path, "ts": int(time.time())}
    else:
        rec[draft_id]["approved"]["0"] = {"name": signer_name, "signature_path": sig_path, "ts": int(time.time())}
    _json_save(REC_FILE, rec)

def _lookup_signature_record(draft_id: str, row_type: str, row_index: int):
    rec = _json_load(REC_FILE, {})
    try:
        return rec[draft_id]["prepared"].get(str(row_index)) if row_type == "prepared" else rec[draft_id]["approved"].get("0")
    except Exception:
        return None

def _get_base_url():
    base = st.secrets.get("APP_BASE_URL","").rstrip("/")
    return base or (st.secrets.get("HTTP_REFERER","") or "").rstrip("/")

def _read_sign_log_records():
    recs = []
    if SIGN_LOG_FILE.exists():
        for line in SIGN_LOG_FILE.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            try:
                recs.append(json.loads(line))
            except Exception:
                pass
    return recs

def _last_rejection_for_draft(draft_id: str):
    recs = _read_sign_log_records()
    for r in reversed(recs):
        if r.get("draft_id") == draft_id and r.get("event") == "approval_rejected":
            return r
    return None

# ---- Tasks & status helpers (use existing TOK/REC files) ----
def _read_tokens():
    return _json_load(TOK_FILE, {})

def _read_records():
    try:
        return json.loads(REC_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _normalize(s):
    return (str(s or "").strip().lower())

def _cfg_mtime(path: str = "config.yaml") -> float:
    try:
        return os.path.getmtime(path)
    except Exception:
        return 0.0
def _normalize_lecturer(raw: dict) -> dict:
    """
    Normalize new schema keys + keep backward compatibility with older keys.
    Output keys used by the app:
      name, email, unit, office, extension, code, designation,
      roles: is_advisor/is_pc/is_cc/is_ec,
      legacy fields: room_no/contact_tel/office_phone/office_hours
    """
    raw = raw or {}
    name = str(raw.get("name", "") or "").strip()
    email = str(raw.get("email", "") or "").strip()

    # New schema
    unit = str(raw.get("unit", "") or "").strip()
    office = str(raw.get("office", "") or "").strip()
    extension = str(raw.get("extension", "") or "").strip()
    designation = str(raw.get("designation", "") or "").strip()
    code = str(raw.get("code", "") or "").strip()

    # Backward compat fallbacks
    room_no = str(raw.get("room_no", "") or "").strip() or office
    contact_tel = (
        str(raw.get("contact_tel", "") or "").strip()
        or str(raw.get("office_phone", "") or "").strip()
        or extension
    )
    office_hours = str(raw.get("office_hours", "") or "").strip()

    roles = {
        "is_advisor": bool(raw.get("is_advisor", False)),
        "is_pc": bool(raw.get("is_pc", False)),
        "is_cc": bool(raw.get("is_cc", False)),
        "is_ec": bool(raw.get("is_ec", False)),
    }

    return {
        "name": name,
        "email": email,
        "unit": unit,
        "office": office,
        "extension": extension,
        "designation": designation,
        "code": code,
        "roles": roles,
        # legacy-style fields your UI already expects
        "room_no": room_no,
        "contact_tel": contact_tel,
        "office_phone": contact_tel,
        "office_hours": office_hours,
    }
def _get_roster(cfg: dict) -> list[dict]:
    roster = cfg.get("lecturers") or []
    if not isinstance(roster, list):
        return []
    out = []
    for it in roster:
        if isinstance(it, dict):
            n = _normalize_lecturer(it)
            if n.get("name") or n.get("email"):
                out.append(n)
    return out
def _email_for_name(name: str) -> str:
    cfg = st.session_state.get("CFG") or load_config(_cfg_mtime())
    roster = _get_roster(cfg)
    n = (name or "").strip().lower()
    for it in roster:
        if (it.get("name") or "").strip().lower() == n:
            return (it.get("email") or "").strip()
    return ""

def _sign_link_for(token: str):
    base = _get_base_url()
    return f"{base}?sign={token}" if base else f"?sign={token}"

def _pending_sign_tasks_for_me():
    """Pending tokens (not used) targeting the signed-in faculty by name/email."""
    prof = st.session_state.get("user_profile") or {}
    me_name  = _normalize(prof.get("name"))
    me_email = _normalize(prof.get("email"))
    toks = _read_tokens()
    out = []
    for tok, info in toks.items():
        if info.get("used_at"):
            continue
        tname  = _normalize(info.get("name"))
        temail = _normalize(info.get("email"))
        if (me_email and temail and me_email == temail) or (me_name and tname and me_name == tname):
            out.append({
                "token": tok,
                "draft_id": info.get("draft_id",""),
                "row_type": info.get("row_type",""),
                "row_index": info.get("row_index", 0),
                "course_code": info.get("course_code",""),
                "course_title": info.get("course_title",""),
                "academic_year": info.get("academic_year",""),
                "semester": info.get("semester",""),
                "link": _sign_link_for(tok),
            })
    return out

def _compute_draft_status(draft_id: str):
    """Return concise status: In progress x/y, Prepared complete, Fully signed."""
    snap = _load_snapshot_if_any(draft_id) or {}
    expected = len(snap.get("prepared_df", []) or [])
    rec = _read_records().get(draft_id, {"prepared":{}, "approved":{}})
    got = len(rec.get("prepared", {}) or {})
    has_approved = bool(rec.get("approved", {}) or {})
    if expected and got >= expected and has_approved:
        return "Fully signed (approved)"
    if expected and got >= expected:
        return "Prepared complete (awaiting approval)"
    if expected:
        return f"In progress ({got}/{expected} prepared)"
    return "No prepared rows"

def _my_issued_links():
    """Tokens issued for drafts I own (_owner_uid) with live status."""
    uid = st.session_state.get("user_code") or ""
    if not uid:
        return []
    toks = _read_tokens()
    rows = []
    for tok, info in toks.items():
        did = info.get("draft_id","")
        snap = _load_snapshot_if_any(did) or {}
        if snap.get("_owner_uid","") != uid:
            continue  # not my draft

        rows.append({
            "token": tok,
            "draft_id": did,
            "row_type": info.get("row_type",""),
            "row_index": info.get("row_index", 0),
            "name": info.get("name",""),
            "course_code": info.get("course_code",""),
            "course_title": info.get("course_title",""),
            "academic_year": info.get("academic_year",""),
            "semester": info.get("semester",""),
            "sections": info.get("sections",""),
            "used_at": info.get("used_at"),
            "note": info.get("note",""),
            "status": _compute_draft_status(did),      # draft-level status
            "token_state": _token_state_label(did, info),  # token-level state
            "link": _sign_link_for(tok),
        })
    return rows

def _token_state_label(draft_id: str, info: dict) -> str:
    """
    Human label for an issued token.
    IMPORTANT: `used_at` means 'closed', not necessarily 'signed'.
    We show ✅ signed only if a signature record exists for that row.
    """
    used_at = info.get("used_at")
    note = (info.get("note") or "").strip().lower()

    # Still open
    if not used_at:
        return "⏳ pending"

    # Closed for reasons other than signing
    if "stale" in note:
        return "⚠️ expired (CDP updated)"
    if "cancel" in note:
        return "🚫 cancelled"
    if "void" in note:
        return "🚫 voided"
    if "reject" in note:
        return "❌ rejected"

    # Closed, but could be signed OR just closed without note.
    row_type  = info.get("row_type", "")
    row_index = int(info.get("row_index", 0) or 0)

    # Verify signature record exists
    sig = _lookup_signature_record(draft_id, row_type, row_index)
    if sig:
        return "✅ signed"

    # Fallback: closed but no signature record
    return "✅ closed"

# ==== end helpers ====

# ==== Draft snapshot storage (top of file, after SIG_DIR/TOK/REC/LOG helpers) ====
DRAFTS_DIR = DATA_DIR / "drafts"
DRAFTS_DIR.mkdir(parents=True, exist_ok=True)

def _current_draft_bundle_dict():
    """Return the CDP bundle as a dict (same data as your sidebar JSON download)."""
    import json as _json
    # reuse your build_bundle() which returns a JSON string
    try:
        return _json.loads(build_bundle())
    except Exception:
        return {}

def _persist_draft_snapshot(draft_id: str) -> Path:
    """Write a snapshot of the current CDP to data/drafts/<draft_id>.json."""
    p = DRAFTS_DIR / f"{draft_id}.json"
    data = _current_draft_bundle_dict()
    # Tag snapshot with the current user (if any). This powers per-user autoload.
    data["_owner_uid"] = st.session_state.get("user_code") or ""
    try:
        p.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass
    return p

def _load_snapshot_if_any(draft_id: str) -> dict | None:
    """Load snapshot dict if exists."""
    p = DRAFTS_DIR / f"{draft_id}.json"
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return None
    return None


st.set_page_config(page_title="UTAS CDP Builder", page_icon="📝", layout="wide")

# ---- Simple "faculty code" login (sidebar) ----
# We’ll accept either an explicit code stored in config.yaml (optional),
# or fall back to the email local-part as the code (before the @).
def _load_cfg_safely():
    try:
        return load_config()
    except Exception:
        return {}

def _build_code_map(cfg: dict) -> tuple[dict, list[dict]]:
    """
    Login map: code -> user_profile.
    Accepts:
      - explicit config 'code'
      - email local-part fallback
    Returns (code_map, collisions)
    """
    roster = _get_roster(cfg)
    code_map = {}
    collisions = []

    def _add(key: str, prof: dict, kind: str):
        if not key:
            return
        if key in code_map and code_map[key].get("email") != prof.get("email"):
            collisions.append({"code": key, "kind": kind, "a": code_map[key], "b": prof})
            # Keep first to avoid hijacking; you can fix collisions in config.yaml
            return
        code_map[key] = prof

    for f in roster:
        name  = (f.get("name") or "").strip()
        email = (f.get("email") or "").strip()
        prof = {
            "name": name,
            "email": email,
            "unit": (f.get("unit") or "").strip(),
            "office": (f.get("office") or "").strip(),
            "extension": (f.get("extension") or "").strip(),
            "designation": (f.get("designation") or "").strip(),
            "roles": f.get("roles") or {},
        }

        explicit = (f.get("code") or "").strip()
        if explicit:
            _add(explicit, prof, "explicit")

        if "@" in email:
            local = email.split("@", 1)[0].strip()
            if local:
                _add(local, prof, "email_localpart")

    return code_map, collisions

@st.cache_data(show_spinner=False)
def load_config(_mtime: float = 0.0) -> dict:
    """
    Cached by config.yaml mtime.
    Returns dict with at least: lecturers (list).
    """
    try:
        with open("config.yaml", "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        if not isinstance(cfg, dict):
            cfg = {}
    except Exception:
        cfg = {}

    cfg.setdefault("lecturers", [])
    if not isinstance(cfg["lecturers"], list):
        cfg["lecturers"] = []
    return cfg

import os
CFG = load_config(_cfg_mtime("config.yaml"))
st.session_state["CFG"] = CFG  # reuse without re-reading
CODE_MAP, CODE_COLLISIONS = _build_code_map(CFG)
def _current_profile() -> dict:
    return st.session_state.get("user_profile") or {}

def _roles() -> dict:
    return (_current_profile().get("roles") or {})

def can_handout_audit() -> bool:
    r = _roles()
    return bool(r.get("is_pc") or r.get("is_cc"))

def is_pc() -> bool:
    return bool(_roles().get("is_pc"))

def is_ec() -> bool:
    return bool(_roles().get("is_ec"))
# --- Signature page router ---
try:
    qp = st.query_params
except Exception:
    qp = st.experimental_get_query_params()
IS_SIGN_LINK = ("sign" in qp)
with st.sidebar:
    if IS_SIGN_LINK:
        st.empty()  # keep sidebar clean for external signers
    else:
        st.markdown("### Faculty Login")

        prev_code = st.session_state.get("user_code", "")
        entered = st.text_input(
            "Enter your code (email local-part or assigned code)",
            value=prev_code,
            type="password"
        )

        if entered:
            if entered in CODE_MAP:
                if entered != prev_code:
                    st.session_state["user_code"]    = entered
                    st.session_state["user_profile"] = CODE_MAP[entered]
                    st.session_state["draft_json_loaded"] = False

                    keep = {"user_code", "user_profile", "SIGN_MODE", "draft_json_loaded", "CFG"}
                    for k in list(st.session_state.keys()):
                        if k not in keep:
                            del st.session_state[k]

                    st.success(f"Signed in as {CODE_MAP[entered]['name']}")
                    st.rerun()
                else:
                    st.success(f"Signed in as {CODE_MAP[entered]['name']}")
            else:
                st.warning("Unknown code. Please check with the Program Coordinator.")

        if st.session_state.get("user_code"):
            prof = _current_profile()
            roles = prof.get("roles") or {}
            st.caption(
                f"**Unit:** {prof.get('unit','-')}\n\n"
                f"**Roles:** "
                + ", ".join([k.replace("is_", "").upper() for k,v in roles.items() if v]) or "-"
            )

# Optional quick "Start fresh" to reset this session’s widgets only
def _clear_fields():
    keep = {"user_code", "user_profile", "SIGN_MODE", "draft_json_loaded"}
    for k in list(st.session_state.keys()):
        if k not in keep:
            del st.session_state[k]
    st.rerun()

with st.sidebar:
    if st.button("Start fresh (clear fields)"):
        _clear_fields()

def _ensure_sched_keys_for_faculty(faculty_list):
    for i, fac in enumerate(faculty_list):
        # per-faculty fields
        st.session_state.setdefault(f"name_{i}",  fac.get("name",""))
        st.session_state.setdefault(f"room_{i}",  fac.get("room_no",""))
        st.session_state.setdefault(f"oh_{i}",    fac.get("office_hours",""))
        st.session_state.setdefault(f"tel_{i}",   fac.get("contact_tel",""))
        st.session_state.setdefault(f"email_{i}", fac.get("email",""))
        # schedule rows
        rows = fac.get("schedule", []) or [{"section":"","day":"","time":"","location":""}]
        st.session_state[f"sched_rows_{i}"] = rows
        for r_i, r in enumerate(rows):
            st.session_state[f"sec_{i}_{r_i}"]  = r.get("section","")
            st.session_state[f"day_{i}_{r_i}"]  = r.get("day","")
            st.session_state[f"time_{i}_{r_i}"] = r.get("time","")
            st.session_state[f"loc_{i}_{r_i}"]  = r.get("location","")
def _normalize_prepared_rows(rows):
    """
    Backward-compatible normalization for Prepared rows.
    Accepts older keys like 'name' and maps them to 'lecturer_name'.
    """
    out = []
    for r in (rows or []):
        if not isinstance(r, dict):
            continue
        out.append({
            "lecturer_name": (r.get("lecturer_name") or r.get("name") or r.get("lecturer") or "").strip(),
            "section_no": (r.get("section_no") or r.get("sections") or r.get("section") or "").strip(),
            "signature": (r.get("signature") or "").strip(),
        })
    return out

def _reset_prepared_widgets():
    """
    Clears widget keys so 'Sync from Faculty' / JSON load can actually refresh defaults.
    """
    prefixes = ("prep_sel_", "prep_sec_", "prep_name_")
    for k in list(st.session_state.keys()):
        if k.startswith(prefixes):
            del st.session_state[k]

def load_draft_into_state(draft):
    st.session_state.setdefault("draft", {})
    st.session_state["draft"]["course"] = draft.get("course", {})
    st.session_state["draft"]["doc"]    = draft.get("doc", {})
    st.session_state["goals_text"]      = draft.get("goals","")
    st.session_state["clos_rows"]       = draft.get("clos_df", [])
    ga = draft.get("graduate_attributes", {}) or {}
    for i in range(1,9):
        st.session_state[f"GA{i}"] = bool(ga.get(f"GA{i}", False))
    srcs = draft.get("sources", {}) or {}
    st.session_state["sources_textbooks"]       = srcs.get("textbooks","")
    st.session_state["sources_reference_books"] = srcs.get("reference_books","")
    st.session_state["sources_e_library"]       = srcs.get("e_library","")
    st.session_state["sources_websites"]        = srcs.get("web_sites","")
    st.session_state["theory_rows"]    = draft.get("theory_df", [])
    st.session_state["practical_rows"] = draft.get("practical_df", [])
    # Seed Assessment split + all four buckets from JSON (overrides prior UI state)
    assess = draft.get("assess", {}) or {}
    st.session_state["draft"]["assess"] = assess

    st.session_state["ass_theory_pct"]    = int(assess.get("theory_pct", 0))
    st.session_state["ass_practical_pct"] = int(assess.get("practical_pct", 0))

    # NEW: buckets
    st.session_state["theory_coursework"]     = list(assess.get("theory_coursework", []))
    st.session_state["theory_final"]          = list(assess.get("theory_final", []))
    st.session_state["practical_coursework"]  = list(assess.get("practical_coursework", []))
    st.session_state["practical_final"]       = list(assess.get("practical_final", []))

    # keep policies if present, but we won't show a UI for them
    st.session_state["draft"]["policies"]  = draft.get("policies", {})

    prep_raw = draft.get("prepared_df", []) or draft.get("prepared_rows", [])
    st.session_state["prepared_rows"] = _normalize_prepared_rows(prep_raw)
    _reset_prepared_widgets()

    st.session_state["date_of_submission"] = draft.get("date_of_submission","")
    ap = draft.get("approved_rows", [])
    if isinstance(ap, list) and ap:
        st.session_state["approved_rows"] = [ap[0]]
    elif isinstance(ap, dict):
        st.session_state["approved_rows"] = [ap]
    else:
        st.session_state["approved_rows"] = [{
            "designation":"Program Coordinator",
            "approved_name":"",
            "approved_date":"",
            "approved_signature":""
        }]
    st.session_state["faculty"] = draft.get("faculty", [])
    _ensure_sched_keys_for_faculty(st.session_state["faculty"])
    st.session_state["draft_json_loaded"] = draft


def _load_latest_snapshot():
    try:
        files = sorted((DRAFTS_DIR.glob("*.json")), key=lambda p: p.stat().st_mtime, reverse=True)
        return json.loads(files[0].read_text(encoding="utf-8")) if files else None
    except Exception:
        return None
def _load_latest_snapshot_for_uid(uid: str):
    """Return the most recent snapshot whose _owner_uid == uid."""
    try:
        files = sorted((DRAFTS_DIR.glob("*.json")), key=lambda p: p.stat().st_mtime, reverse=True)
        for p in files:
            try:
                d = json.loads(p.read_text(encoding="utf-8"))
            except Exception:
                continue
            if d.get("_owner_uid", "") == uid:
                return d
        return None
    except Exception:
        return None

def _load_ai_review_for_token(token: str) -> dict | None:
    try:
        if AI_LOG_FILE.exists():
            for line in AI_LOG_FILE.read_text(encoding="utf-8").splitlines():
                if not line.strip():
                    continue
                try:
                    rec = json.loads(line)
                except Exception:
                    continue
                if rec.get("token") == token and "recommendations_md" in rec:
                    return rec
    except Exception:
        pass
    return None
#for signer page! they review the CDP in a nice layout
GA_LABELS = {
    "GA1": "1. Communication skills",
    "GA2": "2. Teamwork and leadership",
    "GA3": "3. Discipline knowledge and skills",
    "GA4": "4. Creativity and innovation",
    "GA5": "5. Entrepreneurial skills",
    "GA6": "6. Lifelong learning",
    "GA7": "7. Technical and Digital competency",
    "GA8": "8. Critical thinking, analysis, and problem solving",
}
def _render_snapshot_readonly(snap: dict):
    import pandas as pd, re

    d = snap.get("doc", {}) or {}
    c = snap.get("course", {}) or {}

    # ── Course Overview
    st.markdown("#### Course Overview")
    meta_rows = [
        ("Course Code",  c.get("course_code","")),
        ("Course Title", c.get("course_title","")),
        ("Level",        c.get("course_level","")),
        ("Academic Year",d.get("academic_year","")),
        ("Semester",     d.get("semester","")),
        ("Passing Grade",c.get("pass_mark","") or d.get("passing_grade","")),
    ]
    st.dataframe(pd.DataFrame(meta_rows, columns=["Field","Value"]), use_container_width=True)

    # ── Goals
    goals = (snap.get("goals") or "").strip()
    if goals:
        st.markdown("#### Goals")
        st.write(goals)

    # ── CLOs (from clos_df)
    st.markdown("#### Course Learning Outcomes (CLOs)")
    clos = snap.get("clos_df", [])
    if isinstance(clos, list) and clos:
        rows = []
        for i, r in enumerate(clos, start=1):
            rows.append({
                "CLO": f"CLO{i}",
                "Objectives": r.get("objectives",""),
                "Learning Outcomes": r.get("learning_outcomes",""),
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True)
    else:
        st.caption("No CLOs found.")

    # ── Graduate Attributes (course-level selections)
    st.markdown("#### Graduate Attributes (course level)")
    ga_dict = snap.get("graduate_attributes", {}) or {}
    if isinstance(ga_dict, dict) and any(bool(v) for v in ga_dict.values()):
        selected = [k for k,v in ga_dict.items() if v]
        labels   = [GA_LABELS.get(k, k) for k in selected]  # GA_LABELS is defined globally in your app
        st.dataframe(pd.DataFrame(labels, columns=["Selected GA"]), use_container_width=True)
    else:
        st.caption("No Graduate Attributes selected.")

    # ── Sources
    st.markdown("#### Sources")
    srcs = snap.get("sources", {}) or {}
    if srcs:
        src_rows = [
            ("Textbooks",       srcs.get("textbooks","")),
            ("Reference Books", srcs.get("reference_books","")),
            ("E-library",       srcs.get("e_library","")),
            ("Websites",        srcs.get("web_sites","")),
        ]
        st.dataframe(pd.DataFrame(src_rows, columns=["Type","Reference"]), use_container_width=True)
    else:
        st.caption("No sources listed.")

    # ── Weekly Distribution (render both theory & practical)
    def _norm_list(x):
        if isinstance(x, list): return x
        if isinstance(x, str):  return [s.strip() for s in x.split(",") if s.strip()]
        return []

    def _render_weekly(name, rows):
        if isinstance(rows, list) and rows:
            view = []
            for r in rows:
                clos = _norm_list(r.get("clos") or r.get("CLOs") or [])
                gas  = _norm_list(r.get("gas") or r.get("GAs")  or [])
                view.append({
                    "Topic":      r.get("topic",""),
                    "Hours":      r.get("hours",""),
                    "Week":       r.get("week",""),
                    "CLOs":       ", ".join(map(str, clos)),
                    "GAs":        ", ".join(map(str, gas)),
                    "Methods":    r.get("methods",""),
                    "Assessment": r.get("assessment",""),
                })
            st.markdown(f"#### Weekly Distribution — {name}")
            st.dataframe(pd.DataFrame(view), use_container_width=True)

    theory_rows    = snap.get("theory_df", []) or []
    practical_rows = snap.get("practical_df", []) or []
    if not theory_rows and not practical_rows:
        st.markdown("#### Weekly Distribution")
        st.caption("No weekly distribution found.")
    else:
        _render_weekly("Theory", theory_rows)
        _render_weekly("Practical", practical_rows)

    # ── Coverage tables (derived from weekly rows)
    all_weekly = []
    if isinstance(theory_rows, list):    all_weekly += theory_rows
    if isinstance(practical_rows, list): all_weekly += practical_rows

    def _extract_nums(seq):
        out = []
        for s in seq:
            m = re.search(r"(\d+)", str(s))
            if m: out.append(m.group(1))
        return out

    def _coverage(rows, key_label):
        acc = {}  # id -> {"Touchpoints": int, "Total Hours": float}
        for r in rows or []:
            labs = _norm_list(r.get(key_label) or r.get(key_label.upper()) or [])
            ids  = _extract_nums(labs)
            try:
                hrs = float(r.get("hours", 0) or 0)
            except Exception:
                hrs = 0.0
            for i in ids:
                acc.setdefault(i, {"Touchpoints": 0, "Total Hours": 0.0})
                acc[i]["Touchpoints"] += 1
                acc[i]["Total Hours"] += hrs
        table = []
        for i in sorted(acc.keys(), key=lambda x: int(x)):
            label = f"CLO{i}" if key_label.lower() == "clos" else f"GA{i}"
            table.append({"Label": label, **acc[i]})
        return table

    st.markdown("#### Coverage of Learning Outcomes (CLOs)")
    clo_cov = _coverage(all_weekly, "clos")
    if clo_cov:
        st.dataframe(pd.DataFrame(clo_cov), use_container_width=True)
    else:
        st.caption("No CLO coverage table found.")

    st.markdown("#### Coverage of Graduate Attributes (GAs)")
    ga_cov = _coverage(all_weekly, "gas")
    if ga_cov:
        st.dataframe(pd.DataFrame(ga_cov), use_container_width=True)
    else:
        st.caption("No GA coverage table found.")

    # ── Assessment summary (already in your snapshot as a dict)
    assess = snap.get("assess", {}) or {}
    if assess:
        st.markdown("#### Assessment Plan (summary)")
        flat = [{"Component": k, "Value": v} for k, v in assess.items()]
        st.dataframe(pd.DataFrame(flat), use_container_width=True)


if "sign" in qp:
    token = qp["sign"] if isinstance(qp["sign"], str) else qp["sign"][0]
    toks = _json_load(TOK_FILE, {})
    info = toks.get(token)
    st.title("✍️ CDP Digital Signature")

    if not info:
        st.error("Invalid or expired token.")
        st.stop()

    if info.get("used_at"):
        note = str(info.get("note","") or "")
        if "cancel" in note:
            st.info("This sign-off request was cancelled by the CDP creator. Please ignore this link.")
        elif "void" in note:
            st.info("This sign-off request is no longer valid (voided). Please ignore this link.")
        elif "stale" in note:
            st.info("This sign-off request is outdated because the CDP was updated. Please request a new link.")
        else:
            st.info("This token has already been used. Thank you.")
        st.stop()

    # --- NEW: revision guard (CDP changed since link was issued) ---
    expected_rev = (info.get("draft_rev") or "").strip()
    current_rev  = _draft_rev_for_id(info.get("draft_id",""))
    
    if expected_rev and current_rev and expected_rev != current_rev:
        st.warning("⚠️ CDP changed since this request was issued — cancel & re-issue.")
        st.caption("Please contact the CDP creator to void and re-send a fresh sign-off request.")
        _close_token(token, note="stale_revision")  # auto-close so it won't stay pending forever
        st.stop()

    st.write(f"**Signer:** {info.get('name','')}  \n**Draft ID:** {info.get('draft_id','')}  \n**Row:** {info.get('row_type')} #{info.get('row_index')}")
    if info.get("sections"):
        st.caption(f"Sections: {info['sections']}")
    st.session_state["SIGN_MODE"] = True
    # --- Load the frozen CDP snapshot for review ---
    # ... inside: if "sign" in qp:   (after you print Signer/Draft/Row)
    snap = _load_snapshot_if_any(info["draft_id"])
    st.session_state["SIGN_MODE"] = True  # flag, just in case you want it elsewhere
    
    st.markdown("### Review the CDP (read-only)")
    if snap:
        _render_snapshot_readonly(snap)
    else:
        st.warning("No saved CDP snapshot was found for this draft. "
                "Create the sign link from the Sign-off tab (the app will save a snapshot first).")

    
    # --- PD-only AI review: only for 'approved' tokens ---
    if info.get("row_type") == "approved":
        st.markdown("### AI Review (PD only)")
        rec = _load_ai_review_for_token(token)
        if rec and rec.get("recommendations_md"):
            st.markdown(rec["recommendations_md"])
            if rec.get("model"):
                st.caption(f"Model: {rec['model']} • Generated when the approval link was issued.")
        else:
            st.info("No AI review was attached to this approval link.")
    
    # --- PD decision (only for approval links) ---
    decision = "Approve (sign)"
    if info.get("row_type") == "approved":
        decision = st.radio("Decision", ["Approve (sign)", "Reject"], index=0, horizontal=True, key="pd_decision")

        if decision == "Reject":
            reason = st.text_area("Reason for rejection (visible to the CDP creator)", 
                                  key="rej_reason", placeholder="Brief rationale + what to fix…")
            if st.button("❌ Submit rejection", key="btn_submit_rejection"):
                if not reason.strip():
                    st.warning("Please enter a reason for rejection.")
                    st.stop()

                # Audit log the rejection
                _append_sign_log({
                    "ts": int(time.time()),
                    "event": "approval_rejected",
                    "token": token,                         # internal
                    "draft_id": info.get("draft_id",""),
                    "row_type": "approved",
                    "row_index": int(info.get("row_index", 0)),
                    "name": info.get("name",""),
                    "sections": info.get("sections",""),
                    "course_code": info.get("course_code",""),
                    "course_title": info.get("course_title",""),
                    "academic_year": info.get("academic_year",""),
                    "semester": info.get("semester",""),
                    "reason": reason.strip(),
                    "ip": st.secrets.get("REMOTE_IP","")
                })

                # Close the token so it disappears from pending tasks
                _mark_token_used(token)
                # NEW: void all still-open tokens & clear signatures for this draft
                _void_tokens_and_signatures_for_draft(info.get("draft_id",""), reason="voided_on_rejection")
                st.success("Rejection recorded. You may close this window.")
                st.stop()
    if decision != "Approve (sign)":
        st.stop()

    # Signature canvas
    sig = st_canvas(
        fill_color="rgba(0,0,0,0)",
        stroke_width=2,
        stroke_color="#000000",
        background_color="#FFFFFF",
        height=160, width=520, drawing_mode="freedraw",
        key="sign_canvas",
    )

    if st.button("✅ Submit signature"):
        if sig.image_data is None:
            st.warning("Please draw your signature first.")
            st.stop()

        # Convert to transparent PNG where background white → transparent
        arr = sig.image_data.astype("uint8")
        img = Image.fromarray(arr)
        img = img.convert("RGBA")
        data = np.array(img)
        white = (data[:, :, 0:3] == 255).all(axis=2)
        data[white, 3] = 0
        img = Image.fromarray(data, mode="RGBA")

        fname = SIG_DIR / f"{token}.png"
        img.save(str(fname), "PNG")

        # Persist record
        _store_signature_record(
            draft_id=info["draft_id"],
            row_type=info["row_type"],
            row_index=int(info["row_index"]),
            signer_name=info["name"],
            sig_path=str(fname)
        )

        # Audit log
        _append_sign_log({
            "ts": int(time.time()),
            "event": "signature_captured",
            "token": token,                      # internal
            "draft_id": info["draft_id"],        # internal
            "row_type": info["row_type"],        # prepared / approved
            "row_index": int(info["row_index"]),
            "name": info["name"],
            "sections": info.get("sections",""),
            "course_code": info.get("course_code",""),
            "course_title": info.get("course_title",""),
            "academic_year": info.get("academic_year",""),
            "semester": info.get("semester",""),
            "sig_file": str(fname),              # internal
            "ip": st.secrets.get("REMOTE_IP","")
        })


        _mark_token_used(token)
        st.success("Signature saved. You may close this window.")
        # Always halt here so the rest of the editable app (tabs, AI) doesn’t render for signers
        st.stop()
    st.stop()
ALLOWED_LEVELS = ["Bachelor", "Advanced Diploma", "Diploma Second Year", "Diploma First Year"]
SEMESTER_OPTS  = ["Semester I", "Semester II"]


# After reading query params (qp) and before creating tabs:
# Only autoload if a user is "logged in", and load *their* latest snapshot.
if "sign" not in qp and not st.session_state.get("draft_json_loaded"):
    _uid = st.session_state.get("user_code")
    if _uid:  # Logged-in user gets *their* last draft
        latest = _load_latest_snapshot_for_uid(_uid)
        if latest:
            load_draft_into_state(latest)  # seeds all widget keys
            st.session_state["draft_json_loaded"] = True
    # else: no user => do NOT autoload anything (fresh, blank state)

def _stretch_kwargs_for(func):
    try:
        params = inspect.signature(func).parameters
    except Exception:
        params = {}
    if "width" in params:
        return {"width": "stretch"}
    if "use_container_width" in params:
        return {"use_container_width": True}
    return {}
KW_DL  = _stretch_kwargs_for(st.download_button)
KW_BTN = _stretch_kwargs_for(st.button)

import re

def _safe_part(s: str, fallback: str = "") -> str:
    """Filename-safe token: letters/numbers/_ only."""
    s = (str(s or "").strip() or fallback).strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^A-Za-z0-9_]+", "", s)
    return s

def _first_name_from_profile() -> str:
    prof = st.session_state.get("user_profile") or {}
    name = (prof.get("name") or "").strip()
    first = (name.split()[0] if name else "")
    return _safe_part(first, fallback="CDP")

def _academic_year_tag(ay: str) -> str:
    # Examples: "2025-2026" -> "2025", "2025/2026" -> "2025"
    ay = str(ay or "").strip()
    m = re.search(r"(\d{4})", ay)
    return _safe_part(m.group(1) if m else ay, fallback="YYYY")

def _semester_tag(sem: str) -> str:
    sem = str(sem or "").strip().lower()
    if "semester i" in sem or sem.endswith(" i"):
        return "S1"
    if "semester ii" in sem or sem.endswith(" ii"):
        return "S2"
    # fallback: keep something usable
    return _safe_part(sem.upper(), fallback="S?")

def _cdp_download_filename(ext: str = "json") -> str:
    d = st.session_state.get("draft", {}) or {}
    course = (d.get("course") or {})
    doc    = (d.get("doc") or {})

    first = _first_name_from_profile()
    code  = _safe_part(course.get("course_code"), fallback="COURSE")
    year  = _academic_year_tag(doc.get("academic_year"))
    sem   = _semester_tag(doc.get("semester"))

    base = "_".join([p for p in [first, code, year, sem] if p])
    return f"{base}.{ext}"


def normalize_level(s: str) -> str:
    if not s: return ""
    t = s.strip().lower()
    if "advanced" in t and "diploma" in t: return "Advanced Diploma"
    if "second" in t and "diploma" in t: return "Diploma Second Year"
    if "first" in t and "diploma" in t: return "Diploma First Year"
    if "bachelor" in t: return "Bachelor"
    return s if s in ALLOWED_LEVELS else ""

def parse_sections(s: str):
    if not s: return []
    return [p.strip() for p in s.replace("،", ",").split(",") if p.strip()]

def _empty_sched_row(): return {"section":"", "day":"", "time":"", "location":""}
def _empty_clo_row():   return {"objectives":"", "learning_outcomes":""}
def _empty_topic_row():
    return {"topic":"", "hours":1, "week":"", "clos":[], "gas":[], "methods":"", "assessment":""}

def _strip_blank_rows(rows):
    cleaned = []
    for r in rows or []:
        if not isinstance(r, dict): continue
        if "section" in r:
            keys = ("section","day","time","location")
        elif "objectives" in r:
            keys = ("objectives","learning_outcomes")
        elif "designation" in r and "approved_name" in r:
            keys = ("designation","approved_name","approved_date","approved_signature")
        else:
            keys = ("topic","hours","week","methods","assessment")
        rr = {k: str(r.get(k,"")).strip() for k in keys}
        has_text = any(v for v in rr.values())
        if has_text:
            out = dict(r)
            out.update(rr)
            if "clos" in r: out["clos"] = r.get("clos", [])
            if "gas"  in r: out["gas"]  = r.get("gas", [])
            cleaned.append(out)
    return cleaned


@st.cache_data
def load_catalog_from_path(path_str: str):
    p = Path(path_str)
    if not p.exists(): return []
    with open(p, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    items = data.get("courses", data if isinstance(data, list) else [])
    egc = [c for c in items if str(c.get("code","")).upper().startswith("EGC")]
    seen=set(); out=[]
    for c in egc:
        code=str(c.get("code","")).upper().strip()
        if code and code not in seen:
            seen.add(code); out.append(c)
    return out

def course_label(c: dict) -> str:
    code = str(c.get("code","")).strip(); title = str(c.get("title", c.get("name",""))).strip()
    return f"{code} — {title}".strip(" —")

def _parse_hours_from_catalog(c: dict):
    def first_nonempty(d: dict, *keys, cast=int, default=None):
        for k in keys:
            if k in d and d[k] not in (None, "", [], {}):
                try:
                    return cast(d[k])
                except Exception:
                    try:
                        return cast(float(d[k]))
                    except Exception:
                        return d[k]
        return default
    t = first_nonempty(c, "hours_theory_per_week", "hours_theory", "theory_hours",
                          "contact_hours_theory", "lecture_hours", "lec_hours",
                          "theory", "lec", cast=int, default=0)
    p = first_nonempty(c, "hours_practical_per_week", "hours_practical", "practical_hours",
                          "contact_hours_practical", "lab_hours", "pract_hours",
                          "practical", "lab", "labs", cast=int, default=0)
    return int(t or 0), int(p or 0)


def get_roster_names() -> list[str]:
    cfg = st.session_state.get("CFG") or load_config(_cfg_mtime())
    roster = _get_roster(cfg)
    names = [r["name"] for r in roster if r.get("name")]
    return sorted(set(names))


def _subdoc_table(doc, headers, rows, col_widths=None, row_height_in=None):
    sd = doc.new_subdoc()
    t = sd.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    # fixed layout & width
    tblPr = t._tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')
    tblPr.append(tblW)

    for j, h in enumerate(headers):
        t.cell(0, j).text = str(h)

    for r in rows:
        cells = t.add_row().cells
        for j, val in enumerate(r):
            cells[j].text = str(val)

    if col_widths:
        t.autofit = False
        for j, w in enumerate(col_widths):
            try: t.columns[j].width = Inches(w)
            except Exception: pass
        for row in t.rows:
            for j, w in enumerate(col_widths):
                try: row.cells[j].width = Inches(w)
                except Exception: pass

    rh = row_height_in if row_height_in is not None else 0.6
    for idx, row in enumerate(t.rows):
        target = rh*0.75 if idx==0 and rh>0.4 else rh
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(target)
    return sd
def _add_signature_to_cell(cell, img_path, width_inches=1.4):
    try:
        p = cell.paragraphs[0]
        run = p.add_run()
        from docx.shared import Inches
        run.add_picture(img_path, width=Inches(width_inches))
        return True
    except Exception:
        return False

# ----------------------
# PATCH: preload config/catalog up front (unchanged logic)
# ----------------------
cfg = load_config()
catalog = load_catalog_from_path("courses_egc.yaml")
if not catalog:
    catalog = [c for c in (cfg.get("courses") or []) if str(c.get("code","")).upper().startswith("EGC")]

label_to_course = {course_label(c): c for c in catalog}
labels_sorted = ["<manual>"] + sorted(label_to_course.keys())

# Init draft container
if "draft" in st.session_state and not isinstance(st.session_state["draft"], dict):
    st.session_state["draft"] = {}
if "draft" not in st.session_state:
    st.session_state["draft"] = {}

# Sidebar: Save & Load CDP (no template upload — template is shipped with the repo)
if (not IS_SIGN_LINK) and st.session_state.get("user_code"):
    st.sidebar.header("Save & Load CDP")

    # Hide/remove template upload entirely
    uploaded_template = None

    # Upload saved CDP file (internally JSON, but do not expose that term to users)
    st.sidebar.caption("Upload a previously saved CDP file (from this app), then click “Load CDP into app”.")
    json_up = st.sidebar.file_uploader(
        "CDP file",
        type=["json"],
        key="cdp_file_upload",
        label_visibility="collapsed",   # hides the “CDP file” label line
    )
else:
    uploaded_template = None
    json_up = None



# AI Usage & Logs
USAGE_FILE = Path("ai_usage.json")
LOG_FILE = Path("ai_review_logs.jsonl")

def _load_usage():
    try:
        if USAGE_FILE.exists():
            return json.loads(USAGE_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}

def _save_usage(d):
    try:
        USAGE_FILE.write_text(json.dumps(d, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

def _check_and_inc_usage(user_key: str, daily_limit: int = 5):
    usage = _load_usage()
    today = date.today().isoformat()
    cnt = usage.get(user_key, {}).get(today, 0)
    if cnt >= daily_limit:
        return False, cnt
    usage.setdefault(user_key, {})[today] = cnt + 1
    _save_usage(usage)
    return True, cnt + 1
def _reset_usage_today_for(user_key: str | None = None):
    usage = _load_usage()
    today = date.today().isoformat()
    if user_key:
        if user_key in usage and today in usage[user_key]:
            usage[user_key][today] = 0
    else:
        for k in list(usage.keys()):
            if today in usage[k]:
                usage[k][today] = 0
    _save_usage(usage)
def _append_ai_log(record: dict):
    try:
        with AI_LOG_FILE.open("a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    except Exception:
        pass

def _get_faculty_identity():
    # Prefer the logged-in identity
    prof = st.session_state.get("user_profile") or {}
    if prof:
        name  = (prof.get("name")  or "").strip()
        email = (prof.get("email") or "").strip()
        if name or email:
            return (name or "Unknown Faculty"), email

    # Fallback: first faculty row from Tab 1 (legacy behavior)
    fac_list = st.session_state.get("faculty", []) or []
    if fac_list:
        f = fac_list[0]
        name  = (f.get("name")  or "").strip() or "Unknown Faculty"
        email = (f.get("email") or "").strip()
        return name, email

    return "Unknown Faculty", ""

# --- AI review <-> token helpers ---
def _save_ai_review_for_token(token: str, review_md: str, extra: dict | None = None):
    # append to the same AI_LOG_FILE so PD Logs already pick it up
    rec = {
        "ts": datetime.utcnow().isoformat() + "Z",
        "event": "ai_review_auto_on_approval_link",
        "token": token,
        "draft_id": _draft_id(),
        "faculty": _get_faculty_identity()[0],
        "email": _get_faculty_identity()[1],
        "recommendations_md": review_md,
    }
    if isinstance(extra, dict):
        rec.update(extra)
    _append_ai_log(rec)

def _current_user_key():
    code = (st.session_state.get("user_code") or "").strip().lower()
    if code:
        return code
    name, email = _get_faculty_identity()
    if (email or "").strip():
        return email.strip().lower()
    return (name or "unknown").strip().lower()


def _ga_numbers_from_labels(labels):
    # "6. Lifelong learning" -> "6", "GA3" -> "3"
    import re
    out, seen = [], set()
    for s in (labels or []):
        m = re.search(r'(\d+)', str(s))
        if m:
            n = m.group(1)
            if n not in seen:
                seen.add(n)
                out.append(n)
    return out

# AI Prompt
def _build_ai_prompt():
    """Assemble a compact prompt focusing on Weekly Distribution consistency + course-wide coherence."""
    draft = st.session_state.get("draft", {})
    course = draft.get("course", {})
    docinfo = draft.get("doc", {})
    goals   = st.session_state.get("goals_text", "")
    clos    = _strip_blank_rows(st.session_state.get("clos_rows", []))
    # GA course-level indicators: include numbers only
    selected_ga_nums = [str(i) for i in range(1,9) if st.session_state.get(f"GA{i}", False)]

    theory = _strip_blank_rows(st.session_state.get("theory_rows", []))
    practical = _strip_blank_rows(st.session_state.get("practical_rows", []))
    def _clip_text(s: str, limit: int = 800) -> str:
        s = str(s or "")
        return s if len(s) <= limit else (s[:limit] + " …[truncated]")

    def _clip_rows(rows, max_rows=20, max_field=400):
        out = []
        for r in rows[:max_rows]:
            out.append({
                "topic": _clip_text(r.get("topic",""), max_field),
                "hours": r.get("hours",""),
                "week":  _clip_text(r.get("week",""), 50),
                "clos":  list(r.get("clos",[]) or [])[:8],
                "gas":   list(r.get("gas",[]) or [])[:8],
                "methods": _clip_text(r.get("methods",""), max_field),
                "assessment": _clip_text(r.get("assessment",""), max_field),
            })
        return out

    payload = {
        "course": {
            "code": course.get("course_code",""),
            "title": course.get("course_title",""),
            "level": course.get("course_level",""),
            "year":  docinfo.get("academic_year",""),
            "semester": docinfo.get("semester",""),
            "contact_hours": {
                "theory": course.get("hours_theory", 0),
                "practical": course.get("hours_practical", 0),
            },
            "pass_mark": course.get("pass_mark",""),
            "prerequisites": course.get("prerequisite",""),
            "sections": course.get("sections_list", []),
        },
        "goals": _clip_text(goals, 1200),
        "clos": [{"label": f"CLO{i+1}", **row} for i, row in enumerate(clos)],
        "graduate_attributes_course_level": selected_ga_nums,
        "weekly_distribution": {
            "weekly_distribution": {
                "theory": _clip_rows(theory, max_rows=20, max_field=350),
                "practical": _clip_rows(practical, max_rows=20, max_field=350),
            },
        "assessment_split": st.session_state.get("draft", {}).get("assess", {}),
    }}

    system = (
        "You are an academic QA reviewer for Course Delivery Plans (CDPs). "
        "Your job is to check **consistency** and **appropriateness**—especially the Weekly Distribution tables—"
        "against course goals, CLOs and graduate attributes (GA1..GA8). "
        "Be specific, concise, and constructive. Use bullet points. "
        "Flag exaggerations and mismatches (e.g., too many GAs on a single topic, "
        "CLOs not covered, methods/assessment weakly aligned). "
        "Suggest concrete improvements."
    )
    user = (
        "Analyze the following CDP data. Focus on Weekly Distribution coherence with goals/CLOs/GAs, "
        "and also note issues in goals/CLOs if relevant. "
        "Output sections:\n"
        "1) Quick verdict (one short paragraph)\n"
        "2) Issues detected (bullets)\n"
        "3) Suggestions & rewrites (bullets; include adjusted GA/CLO coverage if needed)\n"
        "4) Checklist (pass/fail for: coverage completeness, GA realism, methods/assessment alignment, weekly plan realism)\n"
        f"\nCDP JSON:\n{json.dumps(payload, ensure_ascii=False)}"
    )
    return system, user, payload

def _run_openrouter_review(model: str | None = None, temperature: float = 0.2):
    api_key = st.secrets.get("OPENROUTER_API_KEY")
    if not api_key:
        st.error("OpenRouter API key not set. Add OPENROUTER_API_KEY in Secrets.")
        return None

    system, user, _ = _build_ai_prompt()
    chosen = (model or st.secrets.get("OPENROUTER_DEFAULT_MODEL") or "openrouter/auto").strip()
    FALLBACKS = [ "openrouter/auto", "openai/gpt-4o-mini", "google/gemini-1.5-pro" ]
    tried = []

    def _ascii_header(s: str) -> str:
        return (s or "").encode("ascii", "ignore").decode("ascii")

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "X-Title": "UTAS CDP Builder - AI Review",
    }
    if "HTTP_REFERER" in st.secrets:
        headers["HTTP-Referer"] = _ascii_header(str(st.secrets["HTTP_REFERER"]))

    def _extract_content(data: dict) -> str:
        # OpenAI-format success expected
        if not isinstance(data, dict): return ""
        ch = (data.get("choices") or [])
        if not ch: return ""
        msg = (ch[0] or {}).get("message", {}) or {}
        content = msg.get("content", "")
        # Some providers return a list of blocks
        if isinstance(content, list):
            parts = []
            for part in content:
                if isinstance(part, dict):
                    if "text" in part: parts.append(str(part["text"]))
                    elif "content" in part: parts.append(str(part["content"]))
            content = "\n".join([p for p in parts if p])
        # Anthropic-style refusal sometimes appears separately
        if (not content) and isinstance(msg.get("refusal", ""), str):
            content = f"Refusal: {msg['refusal']}"
        return str(content or "")

    def _call(model_id: str):
        body = {
            "model": model_id,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            "temperature": float(temperature),
            "max_tokens": 1400,
        }
        try:
            resp = requests.post("https://openrouter.ai/api/v1/chat/completions",
                                 headers=headers, json=body, timeout=90)
            text = resp.text
            try:
                data = resp.json()
            except Exception:
                data = {"raw": text}

            meta = {"status": resp.status_code, "model": model_id}
            if not resp.ok:
                err = ""
                if isinstance(data, dict) and "error" in data:
                    err = data["error"].get("message", "")
                if not err: err = text[:500]
                meta["error"] = err
                return False, "", meta

            content = _extract_content(data)
            if not content.strip():
                # Surface helpful slice for PD
                meta["empty_content"] = True
                meta["sample"] = str(data)[:800]
                return True, "", meta  # success but empty
            return True, content, meta

        except Exception as e:
            return False, "", {"error": f"{type(e).__name__}: {e}", "model": model_id}

    # Attempt 1: chosen
    ok, content, meta = _call(chosen); tried.append((chosen, ok, meta))
    # If 400/404 or invalid model → router fallback
    if (not ok) and (meta.get("status") in (400, 404) or "not a valid model id" in str(meta.get("error","")).lower()):
        for fb in FALLBACKS:
            if fb == chosen: continue
            ok, content, meta = _call(fb); tried.append((fb, ok, meta))
            if ok and content.strip(): break
    # If success but empty content → try one more strong fallback
    if ok and not content.strip():
        for fb in FALLBACKS:
            if fb == chosen: continue
            ok2, content2, meta2 = _call(fb); tried.append((fb, ok2, meta2))
            if ok2 and content2.strip():
                ok, content, meta = ok2, content2, meta2
                break

    # PD diagnostics
    if st.session_state.get("PD_MODE"):
        with st.expander("OpenRouter debug (PD only)", expanded=False):
            st.write({"tried": tried})

    if not ok:
        st.error(f"OpenRouter error: {meta.get('error','(no message)')}")
        return None
    if not content.strip():
        st.warning("Model returned an empty message after retries.")
        return ""

    return content


# ----------------------
# PATCH: expand helper to seed ALL faculty widget keys from loaded JSON
# ----------------------


# Load JSON button (only after login)
if (not IS_SIGN_LINK) and st.session_state.get("user_code"):
    if st.sidebar.button("📥 Load CDP into app"):
        import json
        if not json_up:
            st.sidebar.error("Please choose a CDP file first.")
        else:
            try:
                loaded = json.loads(json_up.getvalue().decode("utf-8"))
                load_draft_into_state(loaded)
                st.sidebar.success("CDP loaded.")
                st.rerun()
            except Exception as e:
                st.sidebar.error(f"Could not load the CDP file: {e}")

def build_bundle():
    fac_list = st.session_state.get("faculty", [])
    bundle_fac = []
    for i, f in enumerate(fac_list):
        sched = f.get("schedule", [])
        bundle_fac.append({
            "name": f.get("name", ""),
            "room_no": f.get("room_no", ""),
            "office_hours": f.get("office_hours", ""),
            "contact_tel": f.get("contact_tel", ""),
            "email": f.get("email", ""),
            "schedule": _strip_blank_rows(sched) or [_empty_sched_row()],
        })
    ga_dict = {f"GA{i}": bool(st.session_state.get(f"GA{i}", False)) for i in range(1,9)}
    sources_dict = {
        "textbooks": st.session_state.get("sources_textbooks",""),
        "reference_books": st.session_state.get("sources_reference_books",""),
        "e_library": st.session_state.get("sources_e_library",""),
        "web_sites": st.session_state.get("sources_websites",""),
    }
    bundle = {
        "course": st.session_state.get("draft", {}).get("course", {}),
        "doc": st.session_state.get("draft", {}).get("doc", {}),
        "goals": st.session_state.get("goals_text",""),
        "clos_df": _strip_blank_rows(st.session_state.get("clos_rows", [])),
        "graduate_attributes": ga_dict,
        "sources": sources_dict,
        "theory_df": st.session_state.get("theory_rows", []),
        "practical_df": st.session_state.get("practical_rows", []),
        "assess": st.session_state.get("draft", {}).get("assess", {}),
        "policies": st.session_state.get("draft", {}).get("policies", {}),
        "prepared_df": st.session_state.get("prepared_rows", []),
        "date_of_submission": st.session_state.get("date_of_submission", ""),
        "approved_rows": st.session_state.get("approved_rows", []),
        "faculty": bundle_fac,
    }
    import json
    return json.dumps(bundle, ensure_ascii=False, indent=2)

if (not IS_SIGN_LINK) and st.session_state.get("user_code"):
    # ✅ Give the sidebar JSON download a unique key
    st.sidebar.download_button(
        "💾 Download CDP",
        data=build_bundle(),
        file_name=_cdp_download_filename("json"),
        mime="application/json",
        key="dl_cdp",
        **KW_DL
    )


# App Title ---
st.title("📝 UTAS CDP Builder")


# --- PD Access (shows Tab 8 only to PD) ---
PD_MODE = bool(st.session_state.get("user_code") and is_pc())
if (not IS_SIGN_LINK) and st.session_state.get("user_code") and PD_MODE:
    st.sidebar.success("PC mode enabled")


# --- Gate: require sign-in (or SIGN_MODE) before showing any tabs ---
if not (st.session_state.get("user_code") or st.session_state.get("SIGN_MODE")):
    st.info("Please enter your faculty code in the sidebar to continue.")
    st.stop()
# ---- My tasks (shows for logged-in users) ----
if st.session_state.get("user_code"):
    me = st.session_state.get("user_profile", {})
    if (not IS_SIGN_LINK) and st.session_state.get("user_code"):
        try:
            _expire_stale_tokens_for_draft(_draft_id())
        except Exception:
            pass
    pending = _pending_sign_tasks_for_me()
    issued = _my_issued_links()

    cA, cB = st.columns([1,3])
    with cA:
        st.metric("🔔 Pending sign-offs for you", len(pending))
    with cB:
        with st.expander("My sign-off tasks", expanded=bool(pending)):
            if not pending:
                st.caption("No pending requests.")
            else:
                for t in pending:
                    st.markdown(
                        f"- **{t['course_code']} {t['course_title']}** — {t['semester']} {t['academic_year']} "
                        f"(row: {t['row_type']} #{t['row_index']})  \n"
                        f"  👉 [Open sign page]({t['link']})"
                    )

    with st.expander("Sign-offs I issued"):
        if not issued:
            st.caption("You have not issued any sign-offs yet.")
        else:
            for it in issued:
                label = "Prepared by" if it["row_type"] == "prepared" else "Approved by"
                who   = it.get("name") or f"{it['row_type']} #{it['row_index']}"
                token_state = it.get("token_state", "⏳ pending")
                sections = it.get("sections","") or "-"
                rej = _last_rejection_for_draft(it["draft_id"])
                if rej:
                    when = datetime.fromtimestamp(int(rej.get("ts", 0))).strftime("%Y-%m-%d %H:%M")
                    st.warning(f"❌ Rejected on {when}\n\n**Reason:** {rej.get('reason', '')}")
                st.markdown(
                    f"- **{it['course_code']} {it['course_title']}** — {it['semester']} {it['academic_year']}  \n"
                    f"  {label}: **{who}** • sections: {sections}  \n"
                    f"  Status: **{it['status']}** · {token_state}"
                )

# creating tabs conditionally
# ---- Gate: require sign-in (or SIGN_MODE) before showing any tabs ----
if not (st.session_state.get("user_code") or st.session_state.get("SIGN_MODE")):
    st.info("Please enter your faculty code in the sidebar to continue.")
    st.stop()

TAB_LABELS = [
    "Course & Faculty",
    "Goals, CLOs & Attributes",
    "Sources",
    "Weekly Distribution of the Topics",
    "Assessment Plan",
    "AI Review",
    "Sign-off",
    "Generate",
]

if can_handout_audit():
    TAB_LABELS.append("Handout Audit (PC/CC)")

if PD_MODE:
    TAB_LABELS.append("PD Logs")

_tabs = st.tabs(TAB_LABELS)
TABS = dict(zip(TAB_LABELS, _tabs))

tab1 = TABS["Course & Faculty"]
tab2 = TABS["Goals, CLOs & Attributes"]
tab3 = TABS["Sources"]
tab4 = TABS["Weekly Distribution of the Topics"]
tab5 = TABS["Assessment Plan"]
tab6 = TABS["AI Review"]
tab7 = TABS["Sign-off"]
tab8 = TABS["Generate"]
tab9  = TABS.get("Handout Audit (PC/CC)")
tab10 = TABS.get("PD Logs")


with tab1:
    st.subheader("Course Details")
    choice = st.selectbox("Course (from catalog, or choose <manual>)", labels_sorted, key="course_choice")

    draft_course = st.session_state.get("draft", {}).get("course", {})
    if choice != "<manual>":
        c = label_to_course[choice]
        course_code_val = c.get("code","")
        course_title_val = c.get("title", c.get("name",""))
        t_def, p_def = _parse_hours_from_catalog(c)
        col_hours = st.columns(2)
        with col_hours[0]: hours_theory = st.number_input("Contact Hours — Theory (hrs/week)", 0, 10, int(t_def))
        with col_hours[1]: hours_practical = st.number_input("Contact Hours — Practical (hrs/week)", 0, 10, int(p_def))
        level_default = normalize_level(c.get("course_level",""))
        raw_pr = c.get("prerequisites", [])
        if isinstance(raw_pr, str): prereq_list = [p.strip() for p in raw_pr.split(",") if p.strip()]
        else: prereq_list = [str(x).strip() for x in (raw_pr or []) if str(x).strip()]
    else:
        course_code_val   = st.text_input("Course Code", draft_course.get("course_code",""))
        course_title_val  = st.text_input("Course Name", draft_course.get("course_title",""))
        hours_theory      = st.number_input("Contact Hours — Theory (hrs/week)", 0, 10, int(draft_course.get("hours_theory", 0) or 0))
        hours_practical   = st.number_input("Contact Hours — Practical (hrs/week)", 0, 10, int(draft_course.get("hours_practical", 0) or 0))
        level_default     = normalize_level(draft_course.get("course_level",""))
        _pr               = draft_course.get("prerequisite","")
        prereq_list       = [p.strip() for p in str(_pr).split(",") if p.strip()]

    col_meta = st.columns(3)
    with col_meta[0]: academic_year = st.text_input("Academic Year", st.session_state.get("draft",{}).get("doc",{}).get("academic_year","2025-2026"))
    with col_meta[1]:
        sem_default = st.session_state.get("draft",{}).get("doc",{}).get("semester","Semester I")
        semester = st.selectbox("Semester", SEMESTER_OPTS, index=SEMESTER_OPTS.index(sem_default) if sem_default in SEMESTER_OPTS else 0)
    with col_meta[2]: pass_mark = st.number_input("Passing Grade", 0, 100, int(draft_course.get("pass_mark", 67) or 67))
    course_level = st.selectbox("Course Level", [""] + ALLOWED_LEVELS, index=([""]+ALLOWED_LEVELS).index(level_default) if level_default in ([""]+ALLOWED_LEVELS) else 0)

    ALL_CODES = sorted({str(c.get("code","")) for c in catalog if c.get("code")})
    code_to_label = {c.get("code",""): f"{c.get('code','')} — {c.get('title', c.get('name',''))}" for c in catalog}
    label_to_code = {v:k for k,v in code_to_label.items()}
    options_labels = [code_to_label[c] for c in ALL_CODES]

    ALL_CODES = sorted({str(c.get("code", "")).strip() for c in catalog if str(c.get("code", "")).strip()})
    code_to_label = {
        str(c.get("code", "")).strip(): f"{str(c.get('code','')).strip()} — {c.get('title', c.get('name',''))}"
        for c in catalog
        if str(c.get("code", "")).strip()
    }
    options_labels = [code_to_label[c] for c in ALL_CODES]
    label_to_code = {v: k for k, v in code_to_label.items()}
    
    # ---- Stable prereq widget state keys ----
    PREREQ_LABELS_KEY = "prereq_labels_selected"
    PREREQ_CODES_KEY  = "prereq_codes_store"         # keep your existing store name
    PREREQ_COURSE_PREV_KEY = "prereq_course_prev"
    
    # Decide what we should seed with when the COURSE changes
    # - If catalog provides prereqs: use prereq_list
    # - Else: use whatever is already saved in the draft
    draft_pr_codes = [p.strip() for p in str(draft_course.get("prerequisite", "")).split(",") if p.strip()]
    seed_codes = [str(x).strip() for x in (prereq_list or draft_pr_codes) if str(x).strip()]
    
    # Reset prereq selection ONLY when the selected course changes
    if st.session_state.get(PREREQ_COURSE_PREV_KEY) != choice or PREREQ_LABELS_KEY not in st.session_state:
        st.session_state[PREREQ_CODES_KEY] = list(seed_codes)
    
        seed_labels = []
        for code in seed_codes:
            lbl = code_to_label.get(code)
            if lbl:
                seed_labels.append(lbl)
            else:
                # If something is not in catalog, keep it visible/selectable
                unknown_lbl = f"{code} — (not in catalog)"
                seed_labels.append(unknown_lbl)
                if unknown_lbl not in options_labels:
                    options_labels.append(unknown_lbl)
                label_to_code[unknown_lbl] = code
    
        st.session_state[PREREQ_LABELS_KEY] = seed_labels
        st.session_state[PREREQ_COURSE_PREV_KEY] = choice
    
    # Render with a STABLE key (no versioning)
    st.multiselect(
        "Course Pre-requisite(s) — pick from catalog",
        options_labels,
        key=PREREQ_LABELS_KEY
    )
    
    # Convert selected labels -> codes, persist in store
    selected_labels = st.session_state.get(PREREQ_LABELS_KEY, []) or []
    selected_codes = []
    for lbl in selected_labels:
        code = label_to_code.get(lbl)
        if not code:
            code = (lbl.split(" — ", 1)[0] or "").strip()
        if code:
            selected_codes.append(code)
    
    # Deduplicate while preserving order
    selected_codes = list(dict.fromkeys(selected_codes))
    
    st.session_state[PREREQ_CODES_KEY] = selected_codes
    prereq_final = ", ".join(selected_codes)
    

    sections_str = st.text_input("Section(s) (comma-separated, e.g., 1,4,5,7)", st.session_state.get("draft",{}).get("course",{}).get("sections_str",""))
    sections_list = parse_sections(sections_str)

    _draft = dict(st.session_state.get("draft", {}))
    _course = dict(_draft.get("course", {}))
    _course.update({
        "course_code": course_code_val,
        "course_title": course_title_val,
        "hours_theory": hours_theory,
        "hours_practical": hours_practical,
        "prerequisite": prereq_final,
        "course_level": course_level,
        "pass_mark": pass_mark,
        "sections_str": sections_str,
        "sections_list": sections_list,
    })
    _draft["course"] = _course
    _draft["doc"] = {"academic_year": academic_year, "semester": semester}
    st.session_state["draft"] = _draft
    # right after you set st.session_state["draft"] (Tab 1)
    try:
        _persist_draft_snapshot(_draft_id())
    except Exception:
        pass

    st.markdown("---")
    st.subheader("Faculty Details")
    roster = (load_config().get("lecturers") or [])
    roster_names = [""] + [str(x.get("name","")) for x in roster]

    fac_list = st.session_state.get("faculty", [])
    if not fac_list:
        fac_list = [{
            "name":"", "room_no":"", "office_hours":"", "contact_tel":"", "email":"",
            "schedule":[_empty_sched_row()],
        }]
    st.session_state["faculty"] = fac_list

    addc, delc = st.columns(2)
    with addc:
        if st.button("➕ Add faculty", **KW_BTN):
            fac_list.append({"name":"", "room_no":"", "office_hours":"", "contact_tel":"", "email":"", "schedule":[_empty_sched_row()]})
            st.session_state["faculty"] = fac_list
            st.rerun()
    with delc:
        if st.button("➖ Remove last faculty", **KW_BTN) and len(fac_list) > 1:
            fac_list.pop(); st.session_state["faculty"] = fac_list; st.rerun()

    day_options = ["Sun","Mon","Tue","Wed","Thu","Fri","Sat"]
    for idx, fac in enumerate(fac_list):
        fac_name_for_title = (fac.get("name","") or "").strip() or "Faculty"
        with st.expander(f"Faculty #{idx+1} — {fac_name_for_title}", expanded=True):
            colA, colB = st.columns(2)
            with colA:
                # ----------------------
                # PATCH: only seed from roster when selection CHANGES
                # ----------------------
                sel_key  = f"fac_sel_{idx}"
                prev_key = f"{sel_key}__prev"
                pick = st.selectbox(f"Lecturer/Trainer for Faculty #{idx+1} (pick from roster or leave blank)", roster_names, key=sel_key)

                if st.session_state.get(prev_key, None) != pick:
                    # selection changed -> seed from roster (if any)
                    if pick:
                        info = next((x for x in roster if str(x.get("name","")) == pick), {})
                        st.session_state[f"name_{idx}"]  = info.get("name", fac.get("name",""))
                        # new schema: office -> room_no
                        st.session_state[f"room_{idx}"]  = (
                            info.get("office") or info.get("room_no") or fac.get("room_no","")
                        )
                        st.session_state[f"oh_{idx}"]    = info.get("office_hours", fac.get("office_hours",""))
                        # new schema: extension -> tel
                        st.session_state[f"tel_{idx}"]   = (
                            info.get("extension") or info.get("contact_tel") or info.get("office_phone") or fac.get("contact_tel","")
                        )
                        st.session_state[f"email_{idx}"] = info.get("email", fac.get("email",""))
                    else:
                        # first-time init for manual
                        st.session_state.setdefault(f"name_{idx}",  fac.get("name",""))
                        st.session_state.setdefault(f"room_{idx}",  fac.get("room_no",""))
                        st.session_state.setdefault(f"oh_{idx}",    fac.get("office_hours",""))
                        st.session_state.setdefault(f"tel_{idx}",   fac.get("contact_tel",""))
                        st.session_state.setdefault(f"email_{idx}", fac.get("email",""))
                    st.session_state[prev_key] = pick
                else:
                    # no change; ensure defaults exist (does NOT overwrite typed values)
                    st.session_state.setdefault(f"name_{idx}",  fac.get("name",""))
                    st.session_state.setdefault(f"room_{idx}",  fac.get("room_no",""))
                    st.session_state.setdefault(f"oh_{idx}",    fac.get("office_hours",""))
                    st.session_state.setdefault(f"tel_{idx}",   fac.get("contact_tel",""))
                    st.session_state.setdefault(f"email_{idx}", fac.get("email",""))

                st.text_input("Name of Lecturer", key=f"name_{idx}")
                st.text_input("Lecturer’s Room No.", key=f"room_{idx}")
                st.text_input("Office Hours", key=f"oh_{idx}")  # persists now
                st.text_input("Office Telephone/Extension", key=f"tel_{idx}")
                st.text_input("HCT Email", key=f"email_{idx}")

                fac["name"]         = st.session_state.get(f"name_{idx}", fac.get("name",""))
                fac["room_no"]      = st.session_state.get(f"room_{idx}", fac.get("room_no",""))
                fac["office_hours"] = st.session_state.get(f"oh_{idx}",   fac.get("office_hours",""))
                fac["contact_tel"]  = st.session_state.get(f"tel_{idx}",  fac.get("contact_tel",""))
                fac["email"]        = st.session_state.get(f"email_{idx}",fac.get("email",""))

            with colB:
                st.markdown("**Schedule of the Course Lectures**")
                base_key = f"sched_rows_{idx}"
                if base_key not in st.session_state:
                    rows0 = fac.get("schedule", [_empty_sched_row()])
                    st.session_state[base_key] = rows0
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("➕ Add row", key=f"add_row_{idx}"):
                        st.session_state[base_key].append(_empty_sched_row()); st.rerun()
                with c2:
                    if st.button("➖ Remove row", key=f"del_row_{idx}"):
                        if st.session_state[base_key]: st.session_state[base_key].pop(); st.rerun()
                rows = st.session_state[base_key]
                if not rows: rows = [_empty_sched_row()]; st.session_state[base_key] = rows
                DAY_OPTS = day_options
                for r_i in range(len(rows)):
                    r = rows[r_i]
                    rcols = st.columns(4)
                    with rcols[0]:
                        opts = list(dict.fromkeys((sections_list or [""]) + ([r.get("section","")] if r.get("section","") else [])))
                        st.selectbox("Section", options=opts, index=(opts.index(r.get("section","")) if r.get("section","") in opts else 0), key=f"sec_{idx}_{r_i}")
                    with rcols[1]:
                        st.selectbox("Day", options=DAY_OPTS, index=(DAY_OPTS.index(r.get("day","")) if r.get("day","") in DAY_OPTS else 0), key=f"day_{idx}_{r_i}")
                    with rcols[2]:
                        st.text_input("Time", value=r.get("time",""), key=f"time_{idx}_{r_i}")
                    with rcols[3]:
                        st.text_input("Location", value=r.get("location",""), key=f"loc_{idx}_{r_i}")
                out_rows = []
                for r_i in range(len(st.session_state[base_key])):
                    out_rows.append({"section":st.session_state.get(f"sec_{idx}_{r_i}",""),
                                     "day":st.session_state.get(f"day_{idx}_{r_i}",""),
                                     "time":st.session_state.get(f"time_{idx}_{r_i}",""),
                                     "location":st.session_state.get(f"loc_{idx}_{r_i}","")})
                rows_now = _strip_blank_rows(out_rows) or [_empty_sched_row()]
                fac["schedule"] = rows_now
                if rows_now and any(any(v for v in rr.values()) for rr in rows_now):
                    st.info("Saved schedule rows: " + ", ".join([f"{rr.get('section','')} {rr.get('day','')} {rr.get('time','')} {rr.get('location','')}".strip() for rr in rows_now]))
                else:
                    st.caption("No non-empty schedule rows captured yet.")
            st.session_state["faculty"][idx] = fac

with tab2:
    st.subheader("Goals")
    st.session_state["goals_text"] = st.text_area("Enter course Goals (a short paragraph):", value=st.session_state.get("goals_text",""), height=120)
    st.markdown("---")
    st.subheader("Course Learning Outcomes (CLOs)")
    if not isinstance(st.session_state.get("clos_rows"), list) or not st.session_state["clos_rows"]:
        st.session_state["clos_rows"] = [_empty_clo_row()]
    rows = st.session_state["clos_rows"]
    for idx_row in range(len(rows)):
        r = rows[idx_row]
        c0, c1, c2 = st.columns([0.3, 1.0, 1.0])
        with c0: st.text_input("CLO#", value=f"CLO{idx_row+1}", key=f"clo_label_{idx_row}", disabled=True)
        with c1: st.text_area(f"Objectives (row {idx_row+1})", key=f"clo_obj_{idx_row}", value=r.get("objectives",""), height=100)
        with c2: st.text_area(f"Learning Outcomes (row {idx_row+1})", key=f"clo_out_{idx_row}", value=r.get("learning_outcomes",""), height=100)
        rows[idx_row] = {"objectives": st.session_state.get(f"clo_obj_{idx_row}",""),
                         "learning_outcomes": st.session_state.get(f"clo_out_{idx_row}","")}
    addc, delc = st.columns(2)
    with addc:
        if st.button("➕ Add CLO row"): rows.append(_empty_clo_row()); st.session_state["clos_rows"] = rows; st.rerun()
    with delc:
        if st.button("➖ Remove last CLO row") and len(rows) > 1: rows.pop(); st.session_state["clos_rows"] = rows; st.rerun()
    st.markdown("---")
    st.subheader("Graduate Attributes (tick all that apply)")
    gac1, gac2 = st.columns(2)
    with gac1:
        st.checkbox(GA_LABELS["GA1"], key="GA1"); st.checkbox(GA_LABELS["GA2"], key="GA2"); st.checkbox(GA_LABELS["GA3"], key="GA3"); st.checkbox(GA_LABELS["GA4"], key="GA4")
    with gac2:
        st.checkbox(GA_LABELS["GA5"], key="GA5"); st.checkbox(GA_LABELS["GA6"], key="GA6"); st.checkbox(GA_LABELS["GA7"], key="GA7"); st.checkbox(GA_LABELS["GA8"], key="GA8")
    st.session_state["draft"] = {**st.session_state.get("draft", {}),
        "goals": st.session_state.get("goals_text",""),
        "clos_df": _strip_blank_rows(st.session_state.get("clos_rows", [])),
        "graduate_attributes": {f"GA{i}": bool(st.session_state.get(f"GA{i}", False)) for i in range(1,9)},
    }

with tab3:
    st.subheader("Sources (Title, Author, Publisher, Edition, ISBN no.)")
    st.session_state["sources_textbooks"] = st.text_area("TextBooks", value=st.session_state.get("sources_textbooks",""), height=120)
    st.session_state["sources_reference_books"] = st.text_area("Reference Books", value=st.session_state.get("sources_reference_books",""), height=120)
    st.session_state["sources_e_library"] = st.text_area("E-library reference", value=st.session_state.get("sources_e_library",""), height=100)
    st.session_state["sources_websites"] = st.text_area("Relevant Web Sites", value=st.session_state.get("sources_websites",""), height=100)
    st.session_state["draft"] = {**st.session_state.get("draft", {}),
        "sources": {"textbooks": st.session_state.get("sources_textbooks",""),
                    "reference_books": st.session_state.get("sources_reference_books",""),
                    "e_library": st.session_state.get("sources_e_library",""),
                    "web_sites": st.session_state.get("sources_websites","")}
    }

with tab4:
    st.subheader("Weekly Distribution — Theory & Practical")

    def render_topic_table(kind: str):
        st.markdown(f"**Weekly Distribution {kind} Classes**")
        key_prefix = "theory_rows" if kind == "Theory" else "practical_rows"
        rows = st.session_state.get(key_prefix, [])
        if not rows: rows = [_empty_topic_row()]; st.session_state[key_prefix] = rows
        addc, delc = st.columns(2)
        with addc:
            if st.button(f"➕ Add {kind} row"): rows.append(_empty_topic_row()); st.session_state[key_prefix] = rows; st.rerun()
        with delc:
            if st.button(f"➖ Remove last {kind} row") and len(rows) > 1: rows.pop(); st.session_state[key_prefix] = rows; st.rerun()
        # Build stable base options
        clos_rows_clean = _strip_blank_rows(st.session_state.get("clos_rows", []))
        clo_labels = [f"CLO{i+1}" for i in range(len(clos_rows_clean))] if clos_rows_clean else []

        # Filter GA options to only those checked on Tab 2
        ga_checked_labels = [label for key, label in GA_LABELS.items() if st.session_state.get(key, False)]

        def _merge_options(base, current):
            # Union while preserving order; ensures current selections stay present even if base shrinks
            merged = list(dict.fromkeys(list(base) + list(current or [])))
            return merged

        for i in range(len(rows)):
            r = rows[i]
            # Visual separator + clearer row header
            if i > 0:
                st.markdown("<hr style='border-top:1px solid #ddd; margin:0.75rem 0'/>", unsafe_allow_html=True)
            st.markdown(f"**Row {i+1}**")
            c1, c2, c3 = st.columns([2,1,1])
            with c1: st.text_area("Topics to be covered", key=f"{key_prefix}_topic_{i}", value=r.get("topic",""), height=100)
            with c2:
                hours_opts = list(range(0, 11))  # 0..10
                hours_val = int(r.get("hours", 1))
                idx = hours_val if 0 <= hours_val <= 10 else 0  # index matches value since options = [0..10]
                st.selectbox("Contact Hours", options=hours_opts, index=idx, key=f"{key_prefix}_hours_{i}")

            with c3:st.text_input("Time plan (Week no.)",key=f"{key_prefix}_week_{i}",value=str(r.get("week","")),placeholder="e.g., 2  •  2–3  •  2,4,5")
            c4, c5 = st.columns(2)
            # Stable keys
            clos_key = f"{key_prefix}_clos_{i}"
            gas_key  = f"{key_prefix}_gas_{i}"

            # Seed once (no default=; rely on session_state)
            if clos_key not in st.session_state:
                st.session_state[clos_key] = [x for x in (r.get("clos", []) or []) if x]
            if gas_key not in st.session_state:
                st.session_state[gas_key]  = [x for x in (r.get("gas",  []) or []) if x]

            # Options = filtered base ∪ current selections (prevents clears on rerun)
            clos_opts = _merge_options(clo_labels, st.session_state[clos_key])
            ga_opts   = _merge_options(ga_checked_labels, st.session_state[gas_key])

            with c4:
                st.multiselect("Coverage of Learning Outcomes", options=clos_opts, key=clos_key)
            with c5:
                st.multiselect("Coverage of Graduate Attributes", options=ga_opts, key=gas_key)
            st.text_area("Methods for coverage of Outcomes", key=f"{key_prefix}_methods_{i}", value=r.get("methods",""), height=100)
            st.text_area("Assessment Method(s)/Activities", key=f"{key_prefix}_assessment_{i}", value=r.get("assessment",""), height=100)
            rows[i] = {"topic": st.session_state.get(f"{key_prefix}_topic_{i}",""),
                       "hours": int(st.session_state.get(f"{key_prefix}_hours_{i}", 1)),
                       "week": str(st.session_state.get(f"{key_prefix}_week_{i}", "")).strip(),
                       "clos": list(st.session_state.get(f"{key_prefix}_clos_{i}", [])),
                       "gas": list(st.session_state.get(f"{key_prefix}_gas_{i}", [])),
                       "methods": st.session_state.get(f"{key_prefix}_methods_{i}", ""),
                       "assessment": st.session_state.get(f"{key_prefix}_assessment_{i}", "")}
        st.session_state[key_prefix] = rows

    colA, colB = st.columns(2)
    with colA: render_topic_table("Theory")
    with colB: render_topic_table("Practical")

with tab5:
    st.subheader("Assessment Plan")

    assess_draft = st.session_state.get("draft", {}).get("assess", {}) or {}
    th_def = int(assess_draft.get("theory_pct", 75))
    pr_def = int(assess_draft.get("practical_pct", 25))
    cA, cB = st.columns(2)
    with cA: theory_pct = st.number_input("Theory %", 0, 100, th_def, key="ass_theory_pct")
    with cB: practical_pct = st.number_input("Practical %", 0, 100, pr_def, key="ass_practical_pct")

    def _init_bucket(key, default_rows):
        data = st.session_state.get(key)
        if data is None:
            data = assess_draft.get(key, default_rows)
            st.session_state[key] = data
        return st.session_state[key]

    default_theory_cw = [{"component": "Quizzes", "weight_percent": 10}]
    default_theory_fn = [{"component": "Final Exam", "weight_percent": 30}]
    default_pract_cw  = [{"component": "Experiments", "weight_percent": 15}]
    default_pract_fn  = [{"component": "Practical Test", "weight_percent": 10}]

    theory_coursework = _init_bucket("theory_coursework", default_theory_cw)
    theory_final      = _init_bucket("theory_final", default_theory_fn)
    practical_coursework = _init_bucket("practical_coursework", default_pract_cw)
    practical_final      = _init_bucket("practical_final", default_pract_fn)

    def _row_editor(label, key, rows):
        st.markdown(f"**{label}**")
        addc, delc = st.columns(2)
        with addc:
            if st.button(f"➕ Add row — {label}"):
                rows.append({"component": "", "weight_percent": 0})
                st.session_state[key] = rows
                st.rerun()
        with delc:
            if len(rows) > 0 and st.button(f"➖ Remove last — {label}"):
                rows.pop()
                st.session_state[key] = rows
                st.rerun()
        options = ["Quizzes", "Mid-Sem Exam", "Assignment", "Presentation", "Project",
                   "Experiments", "Practical Test", "Lab Report", "Viva"]
        total = 0
        for i in range(len(rows)):
            c1, c2 = st.columns([3,1])
            comp_key = f"{key}_comp_{i}"
            wt_key   = f"{key}_wt_{i}"
            with c1:
                idx = options.index(rows[i]["component"]) if rows[i]["component"] in options else len(options)
                sel = st.selectbox(f"Component (row {i+1})", options+["(custom)"], index=idx, key=comp_key)
                if sel == "(custom)":
                    rows[i]["component"] = st.text_input(f"Custom component name (row {i+1})",
                                                         value=rows[i]["component"] if rows[i]["component"] not in options else "",
                                                         key=f"{key}_custom_{i}")
                else:
                    rows[i]["component"] = sel
            with c2:
                rows[i]["weight_percent"] = int(st.number_input(f"% (row {i+1})", 0, 100,
                                                                int(rows[i]["weight_percent"]), key=wt_key))
            total += rows[i]["weight_percent"]
        return total

    st.markdown("### Theory")
    c_tc, c_tf = st.columns(2)
    with c_tc:
        tc_sum = _row_editor("Course Work (Theory)", "theory_coursework", theory_coursework)
    with c_tf:
        tf_sum = _row_editor("Final (Theory)", "theory_final", theory_final)

    st.markdown("### Practical")
    c_pc, c_pf = st.columns(2)
    with c_pc:
        pc_sum = _row_editor("Course Work (Practical)", "practical_coursework", practical_coursework)
    with c_pf:
        pf_sum = _row_editor("Final (Practical)", "practical_final", practical_final)

    th_total = tc_sum + tf_sum
    pr_total = pc_sum + pf_sum
    if th_total == st.session_state.get("ass_theory_pct", 0):
        st.success(f"Theory total OK: {th_total}%")
    else:
        st.warning(f"Theory total {th_total}% ≠ Theory split {st.session_state.get('ass_theory_pct',0)}%")
    if pr_total == st.session_state.get("ass_practical_pct", 0):
        st.success(f"Practical total OK: {pr_total}%")
    else:
        st.warning(f"Practical total {pr_total}% ≠ Practical split {st.session_state.get('ass_practical_pct',0)}%")

    # ----------------------
    # PATCH: remove Policies section per cleanup
    # ----------------------
    _draft = dict(st.session_state.get("draft", {}))
    _draft["assess"] = {
        "theory_pct": st.session_state.get("ass_theory_pct", 0),
        "practical_pct": st.session_state.get("ass_practical_pct", 0),
        "theory_coursework": theory_coursework,
        "theory_final": theory_final,
        "practical_coursework": practical_coursework,
        "practical_final": practical_final,
    }
    st.session_state["draft"] = _draft

# =========================
# TAB 6 — AI Review (new)
# =========================
with tab6:
    # AI Button / Review (moved from Generate tab)
    st.markdown("---")
    DEFAULT_MODEL = st.secrets.get("OPENROUTER_DEFAULT_MODEL", "openrouter/auto")
    MODEL_CHOICES = [
        "openrouter/auto",                # safest default (router)
        "anthropic/claude-3.5-sonnet",    # correct ID (no 'openrouter/' prefix)
        "google/gemini-1.5-pro",
        "openai/gpt-4o-mini",
    ]

    st.subheader("AI Review")

    if PD_MODE:
        col_ai1, col_ai2 = st.columns([1, 1])
        with col_ai1:
            ai_model = st.selectbox(
                "Model",
                MODEL_CHOICES,
                index=MODEL_CHOICES.index(DEFAULT_MODEL) if DEFAULT_MODEL in MODEL_CHOICES else 0,
                key="ai_model",
            )
        with col_ai2:
            daily_limit = st.number_input(
                "Daily limit per faculty",
                1,
                20,
                int(st.secrets.get("AI_DAILY_LIMIT", 5)),
                key="ai_daily_limit",
            )
    else:
        ai_model = DEFAULT_MODEL
        daily_limit = int(st.secrets.get("AI_DAILY_LIMIT", 5))

    fac_name, fac_email = _get_faculty_identity()
    user_key = _current_user_key()
    st.caption(f"Counting usage for: **{fac_name or user_key}** {('('+fac_email+')' if fac_email else '')}")

    # Persist PD flag for debug expanders inside functions
    st.session_state["PD_MODE"] = PD_MODE

    if PD_MODE:
        if st.button("♻️ Reset today's AI counter for this user"):
            _reset_usage_today_for(user_key)
            st.success("Today's counter reset for this user.")

    def _peek_usage(user_key: str) -> int:
        usage = _load_usage()
        today = date.today().isoformat()
        try:
            return int(usage.get(user_key, {}).get(today, 0) or 0)
        except Exception:
            return 0

    if st.button("🤖 Run AI Review", **KW_BTN):
        # Use stable key (prefer email, then name)
        user_key = _current_user_key()

        # Show current usage before attempting the run
        used_before = _peek_usage(user_key)
        st.caption(f"AI reviews used today (before this run): {used_before}/{daily_limit} for key: {user_key}")

        # Check and increment usage for today
        allowed, new_cnt = _check_and_inc_usage(user_key, daily_limit=int(daily_limit))

        if not allowed:
            st.warning(f"Daily AI review limit reached for {fac_name or user_key}. Try again tomorrow.")
        else:
            with st.spinner("Running AI review..."):
                ai_text = _run_openrouter_review(model=ai_model)

            if ai_text is None:
                st.error("The OpenRouter call failed. See the error above.")
            elif not str(ai_text).strip():
                st.warning("The AI review returned an empty response. Try changing the model and re-run.")
            else:
                st.success("AI review completed.")
                st.markdown(ai_text)
                st.caption(f"AI reviews used today (after this run): {new_cnt}/{daily_limit}")

                _append_ai_log({
                    "ts": datetime.utcnow().isoformat() + "Z",
                    "faculty": fac_name,
                    "email": fac_email,
                    "course_code": st.session_state.get("draft", {}).get("course", {}).get("course_code",""),
                    "course_title": st.session_state.get("draft", {}).get("course", {}).get("course_title",""),
                    "model": ai_model,
                    "usage_count_today": new_cnt,
                    "recommendations_md": ai_text,
                })


with tab7:
    # top of Tab 7
    if st.button("🔄 Refresh signatures status"):
        st.rerun()
    did = _draft_id()
    expired = _expire_stale_tokens_for_draft(did)
    if expired:
        st.info(f"Auto-closed {expired} stale request(s) (CDP changed since issuance).")
    with st.expander("⚙️ Maintenance / Reset", expanded=False):
        reason = st.text_input("Reason (optional)", key="void_all_reason",
                               placeholder="e.g., major CDP edits; re-sign required")
        confirm = st.checkbox("Confirm void ALL pending requests and clear ALL signatures for this draft",
                              key="void_all_confirm")
    
        if st.button("🚫 Void all sign-offs for this draft", disabled=not confirm):
            did = _draft_id()
            _void_tokens_and_signatures_for_draft(did, reason="voided_by_creator")
            _log_cancel_event(
                event="draft_voided_by_creator",
                draft_id=did,
                row_type="*",
                row_index=-1,
                tokens=[],
                reason=reason,
            )
            st.success("All sign-offs voided and signatures cleared. You can re-issue when ready.")
            st.rerun()

    st.subheader("Sign-off — Prepared & Agreed by")
    st.caption("Seeded from the Faculty schedules on Tab 1. You can edit or add assistants if needed.")

    faculty = st.session_state.get("faculty", [])
    auto_rows = []
    for f in faculty:
        name = (f.get("name") or "").strip()
        secs = []
        for r in (f.get("schedule") or []):
            s = str(r.get("section") or "").strip()
            if s: secs.append(s)
        secs = sorted(set(secs), key=lambda x: (len(x), x))
        if name:
            auto_rows.append({"lecturer_name": name, "section_no": ", ".join(secs), "signature": ""})

    existing = st.session_state.get("prepared_rows", [])
    if not existing:
        existing = auto_rows if auto_rows else [{"lecturer_name":"", "section_no":"", "signature": ""}]
    st.session_state["prepared_rows"] = existing

    names = get_roster_names()
    prep_roster_options = ["— choose —"] + names + ["Other (type manually)"]

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("🔁 Sync from Faculty now"):
            st.session_state["prepared_rows"] = auto_rows if auto_rows else [{"lecturer_name":"", "section_no":"", "signature": ""}]
            _reset_prepared_widgets()
            st.rerun()

    with c2:
        if st.button("➕ Add row"):
            rows = st.session_state.get("prepared_rows", [])
            rows.append({"lecturer_name":"", "section_no":"", "signature": ""})
            st.session_state["prepared_rows"] = rows
            _reset_prepared_widgets()
            st.rerun()
    
    with c3:
        if st.button("➖ Remove last row"):
            rows = st.session_state.get("prepared_rows", [])
            if rows:
                rows.pop()
            st.session_state["prepared_rows"] = rows or [{"lecturer_name":"", "section_no":"", "signature": ""}]
            _reset_prepared_widgets()
            st.rerun()


    rows = st.session_state.get("prepared_rows", [])
    for i in range(len(rows)):
        # OLD:
        # cc1, cc2, cc3 = st.columns([2,1,2])
        # with cc1:  ...  # lecturer select
        # with cc2:  ...  # sections input
        # with cc3:
        #     rows[i]["signature"] = st.text_input(f"Signature (row {i+1})", key=f"prep_sig_{i}", value=rows[i].get("signature",""))
        
        # NEW:
        cc1, cc2 = st.columns([2,1])
        with cc1:
            current_name = (rows[i].get("lecturer_name","") or "").strip()
            # build options that ALWAYS include the current name so it doesn't get wiped
            opts = ["— choose —"] + names
            if current_name and current_name not in opts:
                opts.append(current_name)   # keep it selectable even if not in config.yaml
            opts += ["Other (type manually)"]
            
            default_idx = opts.index(current_name) if current_name in opts else 0
            
            sel = st.selectbox(
                f"Lecturer Name (row {i+1})",
                opts,
                index=default_idx,
                key=f"prep_sel_{i}"
            )
            
            if sel == "— choose —":
                rows[i]["lecturer_name"] = ""
            elif sel == "Other (type manually)":
                rows[i]["lecturer_name"] = st.text_input(
                    f"Type Lecturer Name (row {i+1})",
                    key=f"prep_name_{i}",
                    value=current_name
                )
            else:
                rows[i]["lecturer_name"] = sel

        with cc2:
            rows[i]["section_no"] = st.text_input(f"Section No. (row {i+1})", key=f"prep_sec_{i}", value=rows[i].get("section_no",""))


        # ⬇️ NEW: per-signer link + preview for this Prepared row
        with st.container():
            _di = _draft_id()
            _nm = (rows[i].get("lecturer_name","") or "").strip()
            _secs = (rows[i].get("section_no","") or "").strip()
        
            colL, colR = st.columns([1,3])
            with colL:
                # before issuing, compute who we’re targeting
                _di   = _draft_id()
                _nm   = (rows[i].get("lecturer_name","") or "").strip()
                _mail = _email_for_name(_nm)  # resolves from roster; may be ""
                
                # 1) see if an open token already exists for this exact target
                existing_tok, existing = _find_open_token(
                    draft_id=_di,
                    row_type="prepared",
                    row_index=i,
                    name_or_email=_mail or _nm
                )
                
                if existing_tok:
                    st.info(f"⏳ Pending request already sent to **{_nm or existing.get('name','')}**.")
                    st.caption("If the CDP changes, this link will automatically become invalid and must be re-issued.")
                else:
                    with st.form(f"prep_send_form_{i}", clear_on_submit=False):
                        send = st.form_submit_button(
                            f"📨 Send signature request (row {i+1})",
                            disabled=_busy("prep_busy")
                        )
            
                    if send and not _busy("prep_busy"):
                        _set_busy("prep_busy", True)
                        try:
                            _persist_draft_snapshot(_di)
                            rev = _draft_rev_for_id(_di)
            
                            meta = {
                                "draft_id": _di,
                                "draft_rev": rev,
                                "row_type": "prepared",
                                "row_index": i,
                                "name": _nm,
                                "email": _mail,
                                "sections": _secs,
                                "course_code": st.session_state["draft"]["course"].get("course_code",""),
                                "course_title": st.session_state["draft"]["course"].get("course_title",""),
                                "academic_year": st.session_state["draft"]["doc"].get("academic_year",""),
                                "semester": st.session_state["draft"]["doc"].get("semester",""),
                            }
            
                            tok = _issue_sign_token(meta)
                            ok, err = _email_signoff_request(token=tok, to_name=_nm, to_email=_mail, meta=meta)
                            if ok:
                                st.success("Signature request queued. Email sent.")
                            else:
                                st.success("Signature request queued.")
                                st.warning(f"Email not sent: {err}")
                        finally:
                            _set_busy("prep_busy", False)
                
        
            # 🔕 REMOVE the block that printed the URL:
            # url = st.session_state.get(f"sign_url_prep_{i}")
            # if url:
            #     st.code(url, language="text")
            #     st.caption("Share this link with the lecturer to sign from any device.")
        
            # Signature preview (unchanged)
            rec = _lookup_signature_record(_di, "prepared", i)
            if rec and rec.get("signature_path"):
                st.image(rec["signature_path"], caption="Saved signature", width=220)
            
                with st.form(f"prep_clear_sig_form_{i}", clear_on_submit=False):
                    s_reason = st.text_input("Reason (optional)", key=f"prep_clear_reason_{i}", placeholder="e.g., CDP changed; need fresh sign-off")
                    confirm2 = st.checkbox("Confirm clear saved signature", key=f"prep_clear_confirm_{i}")
                    do_clear = st.form_submit_button("🗑️ Clear signature (require re-sign)", disabled=not confirm2)
            
                if do_clear:
                    ok = _clear_signature_record(_di, "prepared", i, delete_file=False)
                    _log_cancel_event(
                        event="signature_cleared",
                        draft_id=_di,
                        row_type="prepared",
                        row_index=i,
                        name=_nm,
                        email=_mail,
                        sections=(rows[i].get("section_no","") or "").strip(),
                        tokens=[],
                        reason=s_reason,
                    )
                    st.success("Signature cleared. You can send a new request now.")
                    st.rerun()


    st.session_state["prepared_rows"] = rows

    st.text_input("Date of Submission (e.g., 2025-10-01)",
                  key="date_of_submission",
                  value=str(st.session_state.get("date_of_submission", "")))

    st.markdown("---")
    st.subheader("Approved by")
    st.session_state.setdefault("approved_rows", [{
        "designation": "Program Coordinator",
        "approved_name": "",
        "approved_date": "",
        "approved_signature": ""
    }])

    apr = st.session_state["approved_rows"][0]
    desig = st.selectbox("Designation", ["Program Coordinator", "Head of Section", "Head of Department"],
                         index=["Program Coordinator", "Head of Section", "Head of Department"].index(apr.get("designation","Program Coordinator")) if apr.get("designation") in ["Program Coordinator", "Head of Section", "Head of Department"] else 0,
                         key="approved_designation")
    roster_opts = ["— choose —"] + get_roster_names() + ["Other (type manually)"]
    name_sel = st.selectbox("Name", roster_opts, index=0 if not apr.get("approved_name") or apr.get("approved_name") not in roster_opts else roster_opts.index(apr.get("approved_name")), key="approved_name_sel")
    if name_sel == "Other (type manually)":
        name_val = st.text_input("Type Name", value=apr.get("approved_name",""), key="approved_name_manual")
    elif name_sel == "— choose —":
        name_val = ""
    else:
        name_val = name_sel
    date_val = st.text_input("Date", value=apr.get("approved_date",""), key="approved_date")
    #sig_val  = st.text_input("Signature", value=apr.get("approved_signature",""), key="approved_signature")

    st.session_state["approved_rows"] = [{
        "designation": st.session_state.get("approved_designation", "Program Coordinator"),
        "approved_name": name_val,
        "approved_date": date_val,
        "approved_signature": apr.get("approved_signature","")  # keep key but no widget
    }]

    # ⬇️ PD approval — dedupe + AI review (no links shown)
    with st.container():
        # --- Program Director (PD) approval block (keeps existing-token check; adds form+guard) ---
        _di = _draft_id()
        
        # Who is the approver?
        pd_name  = (st.session_state["approved_rows"][0].get("approved_name","") or "").strip()
        pd_email = _email_for_name(pd_name)  # your roster resolver
        
        # Check if an open (unused) approval token already exists for this PD
        existing_tok, existing = _find_open_token(
            draft_id=_di,
            row_type="approved",
            row_index=0,
            name_or_email=pd_email or pd_name
        )
        
        colL, colR = st.columns([1,3])
        with colL:
            if existing_tok:
                st.info(f"Approval request to **{pd_name or existing.get('name','')}** is already pending.")
                st.caption("If the CDP changes, this link will automatically become invalid and must be re-issued.")
            else:
                # Wrap the action in a form so other widget changes don't retrigger this block
                with st.form("pd_submit_form", clear_on_submit=False):
                    submit_pd = st.form_submit_button("📝 Submit for approval", disabled=_busy("ai_busy"))
        
                if submit_pd and not _busy("ai_busy"):
                    _set_busy("ai_busy", True)
                    try:
                        # 1) freeze current CDP snapshot for the approver
                        _persist_draft_snapshot(_di)
                        rev = _draft_rev_for_id(_di)
        
                        # 2) issue approval token
                        meta = {
                            "draft_id": _di,
                            "draft_rev": rev,
                            "row_type": "approved",
                            "row_index": 0,
                            "name": pd_name,
                            "email": pd_email,
                            "sections": "",
                            "course_code": st.session_state["draft"]["course"].get("course_code",""),
                            "course_title": st.session_state["draft"]["course"].get("course_title",""),
                            "academic_year": st.session_state["draft"]["doc"].get("academic_year",""),
                            "semester": st.session_state["draft"]["doc"].get("semester",""),
                        }
                        tok = _issue_sign_token(meta)
                        ok, err = _email_signoff_request(token=tok, to_name=pd_name, to_email=pd_email, meta=meta)
                        if not ok:
                            st.warning(f"Email not sent: {err}")

        
                        # 3) run AI review now and attach it to the token (PD-only)
                        ai_model = (st.session_state.get("ai_model")
                                    or st.secrets.get("OPENROUTER_DEFAULT_MODEL")
                                    or "openrouter/auto")
                        with st.spinner("Running AI review for PD…"):
                            ai_text = _run_openrouter_review(model=ai_model)
        
                        if ai_text and str(ai_text).strip():
                            _save_ai_review_for_token(tok, ai_text, extra={
                                "model": ai_model,
                                "target": {"row_type":"approved","row_index":0,"name":pd_name,"email":pd_email}
                            })
                            st.success("Approval request queued. AI review attached.")
                        else:
                            st.success("Approval request queued (AI review unavailable).")
                    finally:
                        _set_busy("ai_busy", False)
        
        # Optional: show PD signature preview if it exists
        rec = _lookup_signature_record(_di, "approved", 0)
        if rec and rec.get("signature_path"):
            st.image(rec["signature_path"], caption="Saved signature", width=220)
        
            with st.form("pd_clear_sig_form", clear_on_submit=False):
                s_reason = st.text_input("Reason (optional)", key="pd_clear_reason", placeholder="e.g., updated CDP; need fresh approval")
                confirm2 = st.checkbox("Confirm clear saved approval signature", key="pd_clear_confirm")
                do_clear = st.form_submit_button("🗑️ Clear approval signature", disabled=not confirm2)
        
            if do_clear:
                _clear_signature_record(_di, "approved", 0, delete_file=False)
                _log_cancel_event(
                    event="signature_cleared",
                    draft_id=_di,
                    row_type="approved",
                    row_index=0,
                    name=pd_name,
                    email=pd_email,
                    sections="",
                    tokens=[],
                    reason=s_reason,
                )
                st.success("Approval signature cleared. You can submit again when ready.")
                st.rerun()

with tab8:
    st.subheader("Generate")

    # ----------------------
    # PATCH helper: build weekly subdocs
    # ----------------------
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.shared import Length

    def _build_weekly_table(tpl, key_prefix: str, title: str, text_width_emu: int):
        """
        Builds a weekly distribution subdoc table with:
          - Row 1: single merged title cell (bold, centered)
          - Row 2: bold headers with your exact labels
          - Rows 3..N: data
          - Fixed layout + 100% width, column widths proportional to page text width
        """
        import re

        def _ga_labels_to_numbers(labels):
            """Convert GA labels like '6. Lifelong learning' or 'GA6' to '6', keep order, drop dups."""
            out, seen = [], set()
            for s in (labels or []):
                m = re.search(r'(\d+)', str(s))
                if m:
                    n = m.group(1)
                    if n not in seen:
                        seen.add(n)
                        out.append(n)
            return ", ".join(out)

        rows = _strip_blank_rows(st.session_state.get(key_prefix, []))
        if not rows:
            rows = []

        # Compute column widths (in EMUs) from page text width
        # Total weight = 2 + 0.5 + 0.5 + 1 + 1 + 1 + 1 = 7
        base = int(text_width_emu / 7)
        col_widths = [
            2 * base,  # Topics to be covered
            base // 2, # Contact Hours
            base // 2, # Time plan (Week no.)
            base,      # Coverage of Learning Outcomes
            base,      # Coverage of Graduate Attributes
            base,      # Methods for coverage of Outcomes
            base,      # Assessment Method(s)/Activities
        ]

        sd = tpl.new_subdoc()
        table = sd.add_table(rows=2 + max(1, len(rows)), cols=7)
        table.style = "Table Grid"

        # Fixed layout + 100% width
        tblPr = table._tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')  # 100% (pct value is in 1/50ths of a percent)
        tblPr.append(tblW)

        # Row 1: merged title
        title_cell = table.cell(0, 0).merge(table.cell(0, 6))
        p = title_cell.paragraphs[0]
        p.text = ""
        r = p.add_run(title)
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 2: bold headers with exact labels
        headers = [
            "Topics to be covered",
            "Contact Hours",
            "Time plan (Week no.)",
            "Coverage of Learning Outcomes",
            "Coverage of Graduate Attributes",
            "Methods for coverage of Outcomes",
            "Assessment Method(s)/Activities",
        ]
        for j, h in enumerate(headers):
            cell = table.cell(1, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True

        # Apply column widths for all rows
        for row in table.rows:
            for j, w in enumerate(col_widths):
                try:
                    row.cells[j].width = Length(w)
                except Exception:
                    pass

        # Data rows start at index 2
        if rows:
            for i, rdata in enumerate(rows, start=2):
                vals = [
                    rdata.get("topic", ""),
                    rdata.get("hours", ""),
                    rdata.get("week", ""),
                    ", ".join(rdata.get("clos", []) or []),
                    _ga_labels_to_numbers(rdata.get("gas", [])),  # numbers only
                    rdata.get("methods", ""),
                    rdata.get("assessment", ""),
                ]
                for j, v in enumerate(vals):
                    table.cell(i, j).text = str(v)
        else:
            # Keep one empty data row so the table renders with structure
            for j in range(7):
                table.cell(2, j).text = ""

        return sd
    #Footer in geenrated template
    def _footer_line_from_state():
        d = st.session_state.get("draft", {})
        course = d.get("course", {}) or {}
        doc    = d.get("doc", {}) or {}
        dept   = course.get("department", "Engineering Department")
        prog   = course.get("program", "Computer Engineering")
        sem    = (doc.get("semester", "") or "").replace("Semester ", "Sem ")
        year   = doc.get("academic_year", "")
        return f"{dept} / {prog} – {sem} / {year}"
        
    # ✅ Entire generation pipeline runs only when clicked
    if st.button("Generate DOCX", type="primary", **KW_BTN):
        # prefer uploaded file; otherwise fall back to bundled template if present
        if not uploaded_template and Path("Course_Delivery_Plan_Template_placeholders.docx").exists():
            uploaded_template = "Course_Delivery_Plan_Template_placeholders.docx"
        if not uploaded_template:
            st.error("Please upload the official CDP template (.docx) first."); st.stop()
       
        # Build docx template
        tpl = DocxTemplate(uploaded_template)
        # right after you prepare tpl (and before _pr_list / ctx usage)
        draft   = st.session_state.get("draft", {})
        course  = draft.get("course", {})
        docinfo = draft.get("doc", {})
        
        _pr_list = [p.strip() for p in str(course.get("prerequisite","")).split(",") if p.strip()]

        # Faculty schedule tables (build per-faculty subdoc, collect to new_fac)
        fac_list = st.session_state.get("faculty", [])
        new_fac = []
        for f in fac_list:
            rows = _strip_blank_rows(f.get("schedule", []))
            sub = tpl.new_subdoc()
            if rows and any(any(v for v in rr.values()) for rr in rows):
                table = sub.add_table(rows=1+len(rows), cols=4); table.style = "Table Grid"
                hdr = table.rows[0].cells; hdr[0].text = "Section"; hdr[1].text = "Day"; hdr[2].text = "Time"; hdr[3].text = "Location"
                for i, r in enumerate(rows, start=1):
                    cells = table.rows[i].cells
                    cells[0].text = str(r.get("section","")); cells[1].text = str(r.get("day","")); cells[2].text = str(r.get("time","")); cells[3].text = str(r.get("location",""))
            else:
                sub.add_paragraph("No scheduled lectures for this lecturer.")
            f2 = dict(f); f2["schedule_table"] = sub; new_fac.append(f2)

        # Use current template's section to get text width (no need to reload)
        sec = tpl.docx.sections[0]
        text_width = sec.page_width - sec.left_margin - sec.right_margin
        col_label = int(text_width * 0.075)
        remain = int(text_width - col_label)
        half = int(remain // 2)

        # CLOs subdoc
        clos_rows = _strip_blank_rows(st.session_state.get("clos_rows", []))
        clos_sub = tpl.new_subdoc()
        table = clos_sub.add_table(rows=2 + max(1, len(clos_rows)), cols=3); table.style = "Table Grid"
        tblPr = table._tbl.tblPr; tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)
        try:
            table.columns[0].width = Length(col_label)
            table.columns[1].width = Length(half)
            table.columns[2].width = Length(half)
        except Exception: pass
        for row in table.rows:
            try:
                row.cells[0].width = Length(col_label)
                row.cells[1].width = Length(half)
                row.cells[2].width = Length(half)
            except Exception: pass
        h = table.rows[0].cells
        for i, txt in enumerate(["CLO#", "Objectives", "Learning Outcomes"]):
            h[i].text = ""; run = h[i].paragraphs[0].add_run(txt); run.bold = True
        intro = table.rows[1].cells
        intro[0].text = ""
        ir1 = intro[1].paragraphs[0].add_run("This course should enable the students to:"); ir1.bold = True
        ir2 = intro[2].paragraphs[0].add_run("A student who satisfactorily completes the course should be able to:"); ir2.bold = True
        if clos_rows:
            for i, row in enumerate(clos_rows, start=2):
                c = table.rows[i].cells; c[0].text = f"CLO{i-1}"; c[1].text = str(row.get("objectives","")); c[2].text = str(row.get("learning_outcomes",""))
        else:
            c = table.rows[2].cells; c[0].text = "CLO1"; c[1].text = ""; c[2].text = ""

        # GA RichText (bold when selected)
        ga_rt = {}
        for i in range(1,9):
            key = f"GA{i}"; label = GA_LABELS[key]; rt = RichText(); rt.add(label, bold=bool(st.session_state.get(key, False))); ga_rt[f"ga{i}_rt"] = rt

        # Sources block
        sources_sub = tpl.new_subdoc()
        s_table = sources_sub.add_table(rows=5, cols=2); s_table.style = "Table Grid"
        tblPr2 = s_table._tbl.tblPr; tblLayout2 = OxmlElement('w:tblLayout'); tblLayout2.set(qn('w:type'), 'fixed'); tblPr2.append(tblLayout2)
        try:
            s_table.columns[0].width = Length(half); s_table.columns[1].width = Length(half)
        except Exception: pass
        for row in s_table.rows:
            try: row.cells[0].width = Length(half); row.cells[1].width = Length(half)
            except Exception: pass
        hdr_cell = s_table.cell(0,0).merge(s_table.cell(0,1)); hdr_p = hdr_cell.paragraphs[0]; hdr_p.text = ""; hdr_run = hdr_p.add_run("Sources (Title, Author, Publisher, Edition, ISBN no.)"); hdr_run.bold = True
        def set_cell_multiline(cell, text):
            cell.text = ""; lines = str(text or "").splitlines()
            if not lines: cell.paragraphs[0].add_run(""); return
            first=True
            for line in lines:
                if first: cell.paragraphs[0].add_run(line); first=False
                else: p = cell.add_paragraph(""); p.add_run(line)
        labels = ["TextBooks", "Reference Books", "E-library reference", "Relevant Web Sites"]
        vals = [st.session_state.get("sources_textbooks",""), st.session_state.get("sources_reference_books",""), st.session_state.get("sources_e_library",""), st.session_state.get("sources_websites","")]
        for i in range(4):
            s_table.cell(i+1,0).text = labels[i]; set_cell_multiline(s_table.cell(i+1,1), vals[i])

        _pr_list = [p.strip() for p in str(course.get("prerequisite","")).split(",") if p.strip()]

        assess = st.session_state.get("draft", {}).get("assess", {}) or {}
        def _rows_for(bucket):
            arr = assess.get(bucket, []) or []
            return [[r.get("component",""), r.get("weight_percent",0)] for r in arr]
        theory_coursework_table = _subdoc_table(tpl, ["Component","%"], _rows_for("theory_coursework"))
        theory_final_table      = _subdoc_table(tpl, ["Component","%"], _rows_for("theory_final"))
        practical_coursework_table = _subdoc_table(tpl, ["Component","%"], _rows_for("practical_coursework"))
        practical_final_table      = _subdoc_table(tpl, ["Component","%"], _rows_for("practical_final"))

        # Prepared & Approved tables — with embedded signature images
        _di = _draft_id()
        
        # --- Prepared table ---
        prep_rows = st.session_state.get("prepared_rows", []) or []
        prepared_sub = tpl.new_subdoc()
        prep_tbl = prepared_sub.add_table(rows=1 + max(1, len(prep_rows)), cols=4)
        prep_tbl.style = "Table Grid"
        
        # Make table full-width + set fixed layout
        tblPr = prep_tbl._tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)
        tblW = OxmlElement('w:tblW'); tblW.set(qn('w:type'), 'pct'); tblW.set(qn('w:w'), '5000'); tblPr.append(tblW)
        prep_tbl.autofit = False
        
        # Column width plan (approx. % of page text width)
        # 10% | 45% | 15% | 30%  → wider "Lecturer Name" and "Signature"
        total = text_width  # already computed earlier in this handler
        w0 = int(total * 0.10)  # S. No.
        w1 = int(total * 0.45)  # Lecturer Name
        w2 = int(total * 0.15)  # Section No.
        w3 = int(total * 0.30)  # Signature
        
        # Apply widths to all rows (header + data)
        for row in prep_tbl.rows:
            try:
                row.cells[0].width = Length(w0)
                row.cells[1].width = Length(w1)
                row.cells[2].width = Length(w2)
                row.cells[3].width = Length(w3)
            except Exception:
                pass
        
        # Header (bold labels)
        headers = ["S. No.", "Lecturer Name", "Section No.", "Signature"]
        for j, label in enumerate(headers):
            cell = prep_tbl.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(label)
            run.bold = True

        
        if prep_rows:
            for i, r in enumerate(prep_rows, start=1):
                cells = prep_tbl.rows[i].cells
                cells[0].text = str(i)
                cells[1].text = str(r.get("lecturer_name",""))
                cells[2].text = str(r.get("section_no",""))
        
                # signature cell: insert image if available
                sig_cell = cells[3]
                rec = _lookup_signature_record(_di, "prepared", i-1)
                if rec and rec.get("signature_path") and Path(rec["signature_path"]).exists():
                    # clear any existing text
                    for p in list(sig_cell.paragraphs)[1:]:
                        p._element.getparent().remove(p._element)
                    sig_cell.paragraphs[0].clear()
                    if not _add_signature_to_cell(sig_cell, rec["signature_path"], width_inches=1.4):
                        sig_cell.text = r.get("signature","")  # fallback to text
                else:
                    sig_cell.text = r.get("signature","")
        else:
            # keep one empty row so table renders
            cells = prep_tbl.rows[1].cells
            cells[0].text = "1"
            cells[1].text = ""
            cells[2].text = ""
            cells[3].text = ""
        
        prepared_table = prepared_sub  # keep same variable name used in ctx
        
        # --- Approved table ---
        apr_view = (st.session_state.get("approved_rows") or [{}])[0]
        approved_sub = tpl.new_subdoc()
        apr_tbl = approved_sub.add_table(rows=2, cols=4)  # header + 1 data row
        apr_tbl.style = "Table Grid"
        
        # Make table full-width + set fixed layout
        tblPr = apr_tbl._tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed'); tblPr.append(tblLayout)
        tblW = OxmlElement('w:tblW'); tblW.set(qn('w:type'), 'pct'); tblW.set(qn('w:w'), '5000'); tblPr.append(tblW)
        apr_tbl.autofit = False
        
        # Column width plan (Designation, Name wider)
        # 25% | 35% | 15% | 25%
        total = text_width
        w0 = int(total * 0.25)  # Designation
        w1 = int(total * 0.35)  # Name
        w2 = int(total * 0.15)  # Date
        w3 = int(total * 0.25)  # Signature
        
        for row in apr_tbl.rows:
            try:
                row.cells[0].width = Length(w0)
                row.cells[1].width = Length(w1)
                row.cells[2].width = Length(w2)
                row.cells[3].width = Length(w3)
            except Exception:
                pass
        
        # Header (bold labels)
        headers = ["Designation", "Name", "Date", "Signature"]
        for j, label in enumerate(headers):
            cell = apr_tbl.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(label)
            run.bold = True

        
        cells = apr_tbl.rows[1].cells
        cells[0].text = str(apr_view.get("designation",""))
        cells[1].text = str(apr_view.get("approved_name",""))
        cells[2].text = str(apr_view.get("approved_date",""))
        
        sig_cell = cells[3]
        rec_apr = _lookup_signature_record(_di, "approved", 0)
        if rec_apr and rec_apr.get("signature_path") and Path(rec_apr["signature_path"]).exists():
            for p in list(sig_cell.paragraphs)[1:]:
                p._element.getparent().remove(p._element)
            sig_cell.paragraphs[0].clear()
            if not _add_signature_to_cell(sig_cell, rec_apr["signature_path"], width_inches=1.4):
                sig_cell.text = str(apr_view.get("approved_signature",""))
        else:
            sig_cell.text = str(apr_view.get("approved_signature",""))
        
        approved_table = approved_sub  # keep same variable name used in ctx
        

        # Weekly distribution subdocs
        theory_table = _build_weekly_table(
            tpl,
            "theory_rows",
            "Weekly Distribution Theory Classes",
            text_width
        )
        practical_table = _build_weekly_table(
            tpl,
            "practical_rows",
            "Weekly Distribution Practical Classes",
            text_width
        )

        # Context
        ctx = {
            "course_name": course.get("course_title",""),
            "course_code": course.get("course_code",""),
            "hours_theory": course.get("hours_theory", 0),
            "hours_practical": course.get("hours_practical", 0),
            "academic_year": docinfo.get("academic_year",""),
            "semester": docinfo.get("semester",""),
            "passing_grade": course.get("pass_mark",""),
            "course_level": course.get("course_level",""),
            "course_prerequisites": _pr_list,
            "course_prerequisites_str": ", ".join(_pr_list),
            "sections": draft.get("course", {}).get("sections_list", []),

            # text blocks
            "goals": st.session_state.get("goals_text",""),

            # CLOs & GA
            "clos_table": clos_sub,

            # Assessment split + tables
            "assess": {
                "theory_pct": assess.get("theory_pct", 0),
                "practical_pct": assess.get("practical_pct", 0),
            },
            "theory_coursework_table": theory_coursework_table,
            "theory_final_table": theory_final_table,
            "practical_coursework_table": practical_coursework_table,
            "practical_final_table": practical_final_table,

            # Prepared & Approved
            "prepared_table": prepared_table,
            "date_of_submission": st.session_state.get("date_of_submission",""),
            "approved_table": approved_table,

            # Sources
            "sources_table": sources_sub,
            "sources_textbooks":       st.session_state.get("sources_textbooks",""),
            "sources_reference_books": st.session_state.get("sources_reference_books",""),
            "sources_e_library":       st.session_state.get("sources_e_library",""),
            "sources_websites":        st.session_state.get("sources_websites",""),

            # Faculty block (for-loop in template)
            "faculty_list": new_fac,

            # Weekly distribution placeholders
            "theory_table": theory_table,
            "practical_table": practical_table,
        }

        # Add GA RichText placeholders
        ctx.update(ga_rt)

        ctx["footer_line"] = _footer_line_from_state()

        # Render and serve
        tpl.render(ctx)
        out = io.BytesIO()
        tpl.save(out); out.seek(0)
        fname = f"CDP_{ctx.get('course_code','')}_{ctx.get('academic_year','')}_{str(ctx.get('semester','')).replace(' ','_')}.docx"

        # ✅ Unique key for DOCX download
        st.download_button(
            "⬇️ Download DOCX",
            data=out.getvalue(),
            file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_docx_main",
            **KW_DL
        )


def _copy_to_clipboard_button(label: str, text: str, key: str):
    # Streamlit doesn't have native clipboard; use a tiny HTML button.
    safe = json.dumps(text)  # JS-safe string literal

    html = f"""
    <button style="padding:0.4rem 0.7rem;border-radius:8px;border:1px solid #ddd;cursor:pointer;"
            onclick="navigator.clipboard.writeText({safe}); this.innerText='✅ Copied'; setTimeout(()=>this.innerText='{label}', 1200);">
        {label}
    </button>
    """

    kwargs = {"height": 50}
    # Some Streamlit versions don't support 'key' for components.html
    try:
        if "key" in inspect.signature(components.html).parameters:
            kwargs["key"] = key
    except Exception:
        pass

    components.html(html, **kwargs)


@st.cache_data(show_spinner=False)
def _parse_uploaded_cached(file_hash: str, filename: str, ext: str, data: bytes):
    # file_hash is included to ensure cache key stability by content hash.
    suffix = ext.lower()
    p = Path(tempfile.gettempdir()) / f"utas_handout_{file_hash}{suffix}"
    p.write_bytes(data)
    if suffix == ".pdf":
        return parse_pdf(p)
    if suffix == ".pptx":
        return parse_pptx(p)
    if suffix == ".ppt":
        # unsupported; handled by build_corpus() too
        return [{
            "file": filename,
            "unit": "file",
            "heading": "unsupported .ppt",
            "text": "Unsupported format: .ppt (please export/save as .pptx).",
            "captions": [],
        }]
    return [{
        "file": filename,
        "unit": "file",
        "heading": "unsupported",
        "text": f"Unsupported format: {suffix}",
        "captions": [],
    }]
if tab9:
    with tab9:
        st.subheader("Handout Audit (PC/CC)")
        st.caption("Upload current semester handouts (PDF/PPTX). Optionally upload previous semester handouts to assess updates.")

        curr_files = st.file_uploader(
            "Current handouts (ppt/pdf)",
            type=["pdf", "pptx", "ppt"],
            accept_multiple_files=True,
            key="audit_curr_files",
        )
        prev_files = st.file_uploader(
            "Previous semester handouts (optional)",
            type=["pdf", "pptx", "ppt"],
            accept_multiple_files=True,
            key="audit_prev_files",
        )
        lab_cur_up = st.file_uploader(
            "Lab handouts / Practical experiments (optional)",
            type=["pdf", "pptx", "ppt"],
            accept_multiple_files=True,
            key="audit_lab_current",
        )
        lab_prev_up = st.file_uploader(
            "Previous semester lab handouts (optional)",
            type=["pdf", "pptx", "ppt"],
            accept_multiple_files=True,
            key="audit_lab_previous",
        )

        # Persist last result in-session
        if "handout_audit_result" not in st.session_state:
            st.session_state["handout_audit_result"] = None
        if "handout_audit_meta" not in st.session_state:
            st.session_state["handout_audit_meta"] = {}

        def _uploads_to_pairs(files):
            if not files:
                return []
            return [(f.name, f.getvalue()) for f in files]

        with st.form("audit_form", clear_on_submit=False):
            run_btn = st.form_submit_button("🔎 Run Handout Audit", disabled=_busy("ai_busy"))
            if run_btn:
                if not curr_files:
                    st.error("Please upload at least one current handout (PDF/PPTX).")
                else:
                    _set_busy("ai_busy", True)
                    try:
                        # Build corpora FROM BYTES (cached per-file)
                        curr_pairs = _uploads_to_pairs(curr_files)
                        prev_pairs = _uploads_to_pairs(prev_files)
                        lab_cur_pairs = _uploads_to_pairs(lab_cur_up)
                        lab_prev_pairs = _uploads_to_pairs(lab_prev_up)

                        with st.spinner("Parsing handouts (cached)..."):
                            curr_corpus = build_corpus_from_uploads(curr_pairs)

                            prev_corpus = (
                                build_corpus_from_uploads(prev_pairs)
                                if prev_pairs else None
                            )

                            corpus_lab_cur = (
                                build_corpus_from_uploads(lab_cur_pairs)
                                if lab_cur_pairs else None
                            )

                            corpus_lab_prev = (
                                build_corpus_from_uploads(lab_prev_pairs)
                                if lab_prev_pairs else None
                            )

                        # Build CDP bundle from your existing helper
                        cdp_bundle = _current_draft_bundle_dict()

                        payload = prepare_llm_payload(
                            cdp_bundle=cdp_bundle,
                            corpus=curr_corpus,
                            corpus_prev=prev_corpus,
                            corpus_lab_current=corpus_lab_cur,
                            corpus_lab_previous=corpus_lab_prev,
                        )
                        st.session_state["handout_audit_cdp_snapshot"] = payload["cdp_snapshot"]

                        api_key = st.secrets.get("OPENROUTER_API_KEY") or st.secrets.get("openrouter_api_key") or ""
                        app_url = st.secrets.get("APP_URL") if "APP_URL" in st.secrets else None

                        with st.spinner("Auditing handouts with AI..."):
                            audit = run_handout_audit(
                                payload=payload,
                                api_key=api_key,
                                model="anthropic/claude-3.5-sonnet",
                                app_url=app_url,
                                app_title="UTAS CDP Builder",
                                timeout_s=180,
                            )

                        st.session_state["handout_audit_result"] = audit
                        st.session_state["handout_audit_meta"] = {
                            "current_files": [f.name for f in curr_files],
                            "previous_files": [f.name for f in (prev_files or [])],
                            "lab_current_files": [f.name for f in (lab_cur_up or [])],
                            "lab_previous_files": [f.name for f in (lab_prev_up or [])],
                            "ts": int(time.time()),
                        }

                        # Nice-to-have: save into the draft snapshot JSON on disk (local)
                        try:
                            did = _draft_id()
                            snap_path = DRAFTS_DIR / f"{did}.json"
                            if not snap_path.exists():
                                _persist_draft_snapshot(did)
                            snap = json.loads(snap_path.read_text(encoding="utf-8"))
                            snap["handout_audit"] = {
                                "meta": st.session_state["handout_audit_meta"],
                                "result": audit,
                            }
                            snap_path.write_text(json.dumps(snap, indent=2, ensure_ascii=False), encoding="utf-8")
                        except Exception:
                            pass

                        st.success("Handout audit completed.")
                    except Exception as e:
                        st.error(f"Handout audit failed: {e}")
                    finally:
                        _set_busy("ai_busy", False)

        audit = st.session_state.get("handout_audit_result")
        if audit:
            # Render outputs
            pc_tab, cc_tab = st.tabs(["PC Review", "CC Review"])

            with pc_tab:
                rows = audit.get("pc_review", []) or []
                for r in rows:
                    crit = r.get("criterion", "")
                    rating = r.get("rating", "")
                    st.markdown(f"**{crit}** — `{rating}`")
                    st.caption(f"Evidence: {r.get('evidence','')}")
                    st.write(r.get("remarks", ""))

            with cc_tab:
                rows = audit.get("cc_review", []) or []
                for r in rows:
                    crit = r.get("criterion", "")
                    yn = r.get("yes_no", "")
                    st.markdown(f"**{crit}** — `{yn}`")
                    st.caption(f"Evidence: {r.get('evidence','')}")
                    st.write(r.get("remarks", ""))

            st.divider()
            st.subheader("Overall")
            st.write(audit.get("overall_summary", ""))

            if audit.get("action_items"):
                st.subheader("Action Items")
                for a in audit["action_items"]:
                    st.markdown(f"- {a}")

            with st.expander("Trace (evidence highlights)", expanded=False):
                st.json(audit.get("trace", []))

            # Export button
            form_tpl = Path(__file__).parent / "Course_Audit_Form_placeholders.docx"
            bundle = _current_draft_bundle_dict()
            payload = prepare_llm_payload(bundle, corpus=[], corpus_prev=None)

            cdp_snapshot = st.session_state.get("handout_audit_cdp_snapshot") or payload["cdp_snapshot"]
            prof = st.session_state.get("user_profile") or {}
            cdp_snapshot["lecturer_name"] = (prof.get("name") or "").strip()

            docx_bytes = render_course_audit_form_docx(
                template_path=form_tpl,
                audit_json=audit,
                cdp_snapshot=cdp_snapshot,
                # optional (fill if you have them):
                specialization="Computer Engineering",
                unit="CAE",
                pc_name_sign="Dr. Deevyankar",   # or whoever is acting as PC
                cc_member_name_sign="",
                staff_ack_name_sign="",
            )

            st.download_button(
                "📄 Export Audit Summary (.docx)",
                data=docx_bytes,
                file_name="Course_Audit_Form.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_course_audit_form",
                **KW_DL
            )


#Tab 10

if tab10:
    with tab10:
        st.subheader("AI Recommendation Logs")

        records = []
        if AI_LOG_FILE.exists():
            try:
                for line in AI_LOG_FILE.read_text(encoding="utf-8").splitlines():
                    if line.strip():
                        records.append(json.loads(line))
            except Exception as e:
                st.error(f"Failed to read logs: {e}")


        if not records:
            st.info("No AI reviews recorded yet.")
        else:
            import pandas as pd
            df = pd.DataFrame([
                {
                    "Timestamp (UTC)": r.get("ts",""),
                    "Faculty": r.get("faculty",""),
                    "Email": r.get("email",""),
                    "Course": f"{r.get('course_code','')} — {r.get('course_title','')}".strip(" —"),
                    "Model": r.get("model",""),
                    "UsageToday": r.get("usage_count_today",""),
                } for r in records
            ])
            st.dataframe(df, use_container_width=True)

            st.markdown("### Full Recommendations")
            for r in reversed(records):
                with st.expander(f"{r.get('ts','')} — {r.get('faculty','')} — {r.get('course_code','')}", expanded=False):
                    st.markdown(r.get("recommendations_md",""))

            csv_bytes = df.to_csv(index=False).encode("utf-8")
            # ✅ Unique key for log CSV download
            st.download_button(
                "⬇️ Download log (CSV)",
                data=csv_bytes,
                file_name="ai_review_log.csv",
                mime="text/csv",
                key="dl_ai_log_csv",
                **KW_DL
            )
        with st.expander("📊 Course Sign-off Status", expanded=True):
            import pandas as pd, json as _json
        
            # 1) read all signature records (who signed)
            rec = _json_load(REC_FILE, {})
        
            def _read_snapshot(draft_id: str):
                p = (DATA_DIR / "drafts" / f"{draft_id}.json")
                if p.exists():
                    try:
                        return _json.loads(p.read_text(encoding="utf-8"))
                    except Exception:
                        return {}
                return {}
        
            rows = []
            for draft_id, parts in rec.items():
                snap = _read_snapshot(draft_id)
                course = (snap.get("course") or {})
                doc    = (snap.get("doc") or {})
                prepared_total = len(snap.get("prepared_df") or snap.get("prepared_rows") or [])
                prepared_signed = len((parts.get("prepared") or {}).keys())
                approved_signed = bool((parts.get("approved") or {}).get("0"))
        
                if prepared_total == 0 and not approved_signed and prepared_signed == 0:
                    status = "Not started"
                elif prepared_signed < max(1, prepared_total):
                    status = "In progress"
                elif prepared_signed >= max(1, prepared_total) and not approved_signed:
                    status = "Waiting for approval"
                else:
                    status = "Approved"
                # after you compute prepared_signed / prepared_total / approved_signed:
                rej = _last_rejection_for_draft(draft_id)
                if rej:
                    status = "Rejected — needs revision"
                elif prepared_signed < max(1, prepared_total):
                    status = "In progress"
                elif prepared_signed >= max(1, prepared_total) and not approved_signed:
                    status = "Waiting for approval"
                else:
                    status = "Approved"
        
                rows.append({
                    "Course": f"{course.get('course_code','')} — {course.get('course_title','')}".strip(" —"),
                    "AY": doc.get("academic_year",""),
                    "Semester": doc.get("semester",""),
                    "Prepared (signed/total)": f"{prepared_signed}/{prepared_total}",
                    "Approved?": "Yes" if approved_signed else "No",
                    "Status": status,
                    "Draft ID": draft_id,  # keep for reference/debug
                })
        
            if not rows:
                st.info("No course records yet.")
            else:
                df = pd.DataFrame(rows).sort_values(["Status","Course"])
                st.dataframe(df, use_container_width=True)

        
        #sign off logs in pd tab
        
        with st.expander("🗂️ Sign-off Audit Log (PD)", expanded=False):
            import pandas as pd, datetime as _dt, json as _json
        
            rows = []
            if SIGN_LOG_FILE.exists():
                for line in SIGN_LOG_FILE.read_text(encoding="utf-8").splitlines():
                    if not line.strip(): continue
                    rec = _json.loads(line)
                    rows.append(rec)
        
            if not rows:
                st.info("No sign-off records yet.")
            else:
                def _fmt_ts(ts):
                    try:
                        return _dt.datetime.fromtimestamp(int(ts)).strftime("%Y-%m-%d %H:%M")
                    except Exception:
                        return ""
        
                nice = []
                for r in rows:
                    nice.append({
                        "When": _fmt_ts(r.get("ts")),
                        "Type": r.get("row_type", "").title(),                    # Prepared / Approved
                        "Signer": r.get("name",""),
                        "Sections": r.get("sections",""),
                        "Course": f"{r.get('course_code','')} — {r.get('course_title','')}".strip(" —"),
                        "AY": r.get("academic_year",""),
                        "Semester": r.get("semester",""),
                        # keep internals out of the main table:
                        # "_token": r.get("token",""), "_draft_id": r.get("draft_id",""),
                    })
        
                df = pd.DataFrame(nice)
                st.dataframe(df, use_container_width=True)
        
                # Optional: CSV download
                st.download_button(
                    "⬇️ Download sign-off log (CSV)",
                    data=df.to_csv(index=False).encode("utf-8"),
                    file_name="signoff_log_clean.csv",
                    mime="text/csv",
                    key="dl_signoff_log_csv",
                    **KW_DL
                )
